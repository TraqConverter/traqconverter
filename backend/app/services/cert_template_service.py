"""Certification template engine.

A cert template is just a DOCX (or PDF) that the user uploads to the
Certifications library. The DOCX can contain placeholder tokens of
the form ``{{field_name}}`` anywhere in its text — headers, body
paragraphs, table cells, signature lines, footers — and at export
time we substitute every recognised token for the corresponding
value drawn from the project, user, and team.

Why placeholders are great here
-------------------------------
- The user controls the visual design (their letterhead, ISO 17100
  stamp graphic, signature block, language) — we don't have to
  hand-craft a template per customer.
- The placeholder set is finite + documented (see SUPPORTED_FIELDS
  below) so the UI can show "we found 6 known placeholders + 1
  unknown" right after upload.
- Substitution is XML-level: we don't re-render or re-parse the
  document, so styles, images, tables, signatures, and any embedded
  fonts come through untouched.

Format support
--------------
Only DOCX templates are supported. PDFs are flat — we'd need to do
either (a) form-field substitution, which requires the user to
build the PDF as a fillable form, or (b) re-render which destroys
the original layout. DOCX gives us both.
"""
from __future__ import annotations

import io
import logging
import re
import uuid
from datetime import datetime
from typing import Any, Iterable

logger = logging.getLogger(__name__)


# Every placeholder we know how to substitute. Order matters for the
# UI: this controls the order documented to the user when they upload
# a template.
SUPPORTED_FIELDS: tuple[tuple[str, str], ...] = (
    ("translator_name",     "Logged-in user's full name (falls back to email username)"),
    ("translator_email",    "Logged-in user's email address"),
    ("translator_title",    "Logged-in user's role (e.g. Project manager)"),
    ("date",                "Today's date in ISO format (2026-05-27)"),
    ("date_long",           "Today's date in long format (27 May 2026)"),
    ("source_language",     "Project's source language"),
    ("target_language",     "Project's target language"),
    ("document_name",       "Project's source file name"),
    ("page_count",          "Number of pages in the source document"),
    ("word_count",          "Total source words across all segments"),
    ("segment_count",       "Total segment count in the project"),
    ("project_id",          "Project's unique identifier"),
    ("certificate_number",  "Auto-generated unique certificate number"),
    ("team_name",           "Team / company name"),
    ("company_address",     "Team / company address (from Settings)"),
)
SUPPORTED_FIELD_NAMES: frozenset[str] = frozenset(name for name, _ in SUPPORTED_FIELDS)

# Pattern matches ``{{anything_alpha_or_underscore_or_digit}}`` with
# optional surrounding spaces inside the braces, e.g. both
# ``{{date}}`` and ``{{ date }}`` work.
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


# ============================================================
# Token discovery
# ============================================================


def scan_docx_for_tokens(docx_bytes: bytes) -> dict[str, Any]:
    """Open a DOCX in-memory and return every placeholder we can see.

    The return shape is intentionally explicit so the upload UI can
    render it directly:

        {
          "found": ["translator_name", "date", "source_language"],
          "unknown": ["client_dispatcher"],     # placeholders we don't fill
          "supported": [{"name": ..., "description": ...}, ...]
        }

    Unknown placeholders are surfaced (not dropped) so the user knows
    their template references something that won't get substituted.
    """
    found: list[str] = []
    try:
        from docx import Document  # python-docx

        doc = Document(io.BytesIO(docx_bytes))
        for text in _iter_docx_text(doc):
            for match in _PLACEHOLDER_RE.finditer(text):
                token = match.group(1)
                if token not in found:
                    found.append(token)
    except Exception as e:
        logger.warning("Cert template scan failed: %s", e)

    known = [t for t in found if t in SUPPORTED_FIELD_NAMES]
    unknown = [t for t in found if t not in SUPPORTED_FIELD_NAMES]
    return {
        "found": known,
        "unknown": unknown,
        "supported": [
            {"name": name, "description": desc}
            for name, desc in SUPPORTED_FIELDS
        ],
    }


def _iter_docx_text(doc) -> Iterable[str]:
    """Walk every text container in a python-docx Document and yield
    the visible text. Covers paragraphs, table cells, headers, and
    footers."""
    for para in doc.paragraphs:
        yield para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para.text
    for section in doc.sections:
        for container in (section.header, section.footer):
            if container is None:
                continue
            for para in container.paragraphs:
                yield para.text
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para.text


# ============================================================
# Value building
# ============================================================


def build_substitution_values(
    *,
    user: Any,
    project: Any,
    team: Any = None,
    extra: dict[str, str] | None = None,
) -> dict[str, str]:
    """Construct the {token: value} map applied during substitution.

    Caller can pass an `extra` dict to override or add fields (e.g. a
    pre-generated certificate_number that should match what's written
    to the audit log).
    """
    now = datetime.utcnow()

    def _safe(val: Any) -> str:
        if val is None:
            return ""
        return str(val)

    full_name = _safe(getattr(user, "full_name", "")).strip() or _safe(
        getattr(user, "email", "")
    ).split("@")[0]
    role = _safe(getattr(user, "role", "")).strip() or "Translator"

    values: dict[str, str] = {
        "translator_name": full_name,
        "translator_email": _safe(getattr(user, "email", "")),
        "translator_title": role,
        "date": now.strftime("%Y-%m-%d"),
        "date_long": now.strftime("%d %B %Y"),
        "source_language": _safe(getattr(project, "source_language", "")),
        "target_language": _safe(getattr(project, "target_language", "")),
        "document_name": _safe(getattr(project, "file_name", "")),
        "page_count": _safe(getattr(project, "page_count", 0)),
        "word_count": _safe(
            getattr(project, "word_count", None)
            or getattr(project, "total_word_count", "")
        ),
        "segment_count": _safe(getattr(project, "total_segments", 0)),
        "project_id": _safe(getattr(project, "id", "")),
        "certificate_number": _make_certificate_number(project),
        "team_name": _safe(getattr(team, "name", "")) if team else "",
        "company_address": _safe(getattr(team, "address", "")) if team else "",
    }
    if extra:
        values.update({k: _safe(v) for k, v in extra.items()})
    return values


def _make_certificate_number(project: Any) -> str:
    """Produce a stable, audit-friendly certificate number.

    Format: ``CERT-YYYY-<8-char-project-uuid-prefix>`` — gives every
    project a unique, year-scoped number that's also easy to look up
    later. Stable per project, so re-exporting doesn't churn it.
    """
    year = datetime.utcnow().year
    pid = getattr(project, "id", None)
    prefix = ""
    if isinstance(pid, uuid.UUID):
        prefix = pid.hex[:8].upper()
    elif pid:
        prefix = str(pid).replace("-", "")[:8].upper()
    if not prefix:
        prefix = uuid.uuid4().hex[:8].upper()
    return f"CERT-{year}-{prefix}"


# ============================================================
# Substitution
# ============================================================


def substitute_in_docx(
    docx_bytes: bytes,
    values: dict[str, str],
) -> bytes:
    """Return the DOCX bytes with every ``{{token}}`` replaced by the
    matching value. Unknown tokens are left in place so the user can
    spot them if a substitution is missing.

    Substitution is run-aware: Word often splits a single visible
    placeholder ``{{date}}`` across multiple ``<w:r>`` runs when the
    user touches it during editing. We join the runs within a
    paragraph, do the substitution, then write the result back into
    the first run and clear the others — this keeps the paragraph's
    paragraph-level formatting (alignment, style) intact while
    avoiding the "placeholder doesn't get replaced because it's
    fragmented across runs" issue you'd hit with a naive
    ``run.text.replace`` loop.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))

    for para in doc.paragraphs:
        _substitute_paragraph(para, values)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _substitute_paragraph(para, values)
    for section in doc.sections:
        for container in (section.header, section.footer):
            if container is None:
                continue
            for para in container.paragraphs:
                _substitute_paragraph(para, values)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _substitute_paragraph(para, values)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _substitute_paragraph(paragraph, values: dict[str, str]) -> None:
    """Replace every supported placeholder in a paragraph, even when
    Word has split the placeholder across multiple runs."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    def _replace(match: re.Match) -> str:
        token = match.group(1)
        if token in values:
            return values[token]
        # Leave unknown placeholders untouched so the user can spot them.
        return match.group(0)

    new_text = _PLACEHOLDER_RE.sub(_replace, full_text)
    if new_text == full_text:
        return

    runs = paragraph.runs
    if not runs:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
