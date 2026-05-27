"""Export service.

The translation processor rebuilds a layout-preserving copy of the source
file and uploads it to S3 (`project.output_file`). For exports we prefer
that rebuilt file because it visually resembles the original — table
boxes, ID-card layouts and certificate borders are preserved.

Only when the rebuild is missing (older projects, or a rebuild that
failed mid-flight) do we fall back to a flat paragraph stack.
"""
from io import BytesIO
from datetime import datetime
import logging

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

logger = logging.getLogger(__name__)


# ==============================
# CERTIFICATION BLOCK
# ==============================

def build_certification(user_email: str):
    return [
        "CERTIFIED TRANSLATION",
        "",
        "I hereby certify that this translation is accurate and complete.",
        "",
        f"Translator: {user_email}",
        f"Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        "",
        "----------------------------------------",
        "",
    ]


# ==============================
# Try to load the layout-preserving DOCX rebuilt by the worker.
# ==============================

def _try_layout_docx_from_project(project) -> BytesIO | None:
    """Return the rebuilt DOCX as a BytesIO, prepending the certification
    page. Returns None if there's no rebuilt file or it isn't a DOCX."""
    if project is None:
        return None
    output_key = getattr(project, "output_file", None)
    if not output_key:
        return None
    if getattr(project, "source_kind", "").upper() != "DOCX":
        return None

    try:
        from app.services.s3_service import generate_presigned_download_url
        import requests
        url = generate_presigned_download_url(output_key)
        r = requests.get(url, timeout=15)
        r.raise_for_status()
    except Exception as e:
        logger.warning(f"Couldn't fetch rebuilt DOCX from S3: {e}")
        return None

    try:
        rebuilt = Document(BytesIO(r.content))
    except Exception as e:
        logger.warning(f"Rebuilt DOCX wouldn't open: {e}")
        return None

    out = Document()
    # Document body first.
    for para in rebuilt.paragraphs:
        new_p = out.add_paragraph()
        new_p.style = para.style
        for run in para.runs:
            r = new_p.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            if run.font and run.font.size:
                r.font.size = run.font.size
    for table in rebuilt.tables:
        new_table = out.add_table(rows=len(table.rows), cols=len(table.columns))
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                new_table.cell(ri, ci).text = cell.text

    # Certification block appended at the end.
    out.add_paragraph("")
    for line in build_certification(getattr(project, "_export_user_email", "")):
        out.add_paragraph(line)

    buffer = BytesIO()
    out.save(buffer)
    buffer.seek(0)
    return buffer


def _build_layout_pdf_live(segments, project):
    """Build the export PDF in the new STRUCTURED format:

        Page 1+      : original document (image or PDF) verbatim
        Pages N+     : clean structured translation
        Final page   : certification block (optional company logo)

    Replaces the old in-place overlay (which produced bilingual residue
    on image documents) with a professionally-typeset translation that
    preserves alignment, bold/italic emphasis, placeholders for
    non-text elements, and source page breaks.
    """
    if project is None or not segments:
        return None
    kind = (getattr(project, "source_kind", "") or "").upper()
    if kind not in ("PDF", "IMAGE"):
        return None

    file_key = getattr(project, "file_path", None)
    if not file_key:
        return None

    import os
    import tempfile
    from datetime import datetime
    from pathlib import Path
    from app.services.s3_service import generate_presigned_download_url
    from app.services.structured_translation_renderer import (
        render_structured_export,
    )
    import requests

    try:
        url = generate_presigned_download_url(file_key)
        r = requests.get(url, timeout=20)
        r.raise_for_status()
    except Exception as e:
        logger.warning(f"Couldn't fetch source from S3: {e}")
        return None

    tmp_dir = Path(tempfile.mkdtemp(prefix="traqexport_"))
    try:
        ext = (
            ".pdf"
            if kind == "PDF"
            else os.path.splitext(file_key)[1] or ".jpg"
        )
        src_path = tmp_dir / f"source{ext}"
        with open(src_path, "wb") as f:
            f.write(r.content)

        # Build (translated, layout) pairs from current segments.
        pairs = []
        for s in segments:
            if not s.translated_text or not s.translated_text.strip():
                continue
            meta = dict(s.layout_meta or {})
            meta["source_text"] = s.source_text
            pairs.append((s.translated_text, meta))

        if not pairs:
            return None

        # Resolve a company logo: per-user S3 logo first, then optional
        # global default from settings.
        logo_path = None
        user_logo_key = getattr(project, "_export_user_logo_key", None)
        if user_logo_key:
            try:
                from app.services.s3_service import (
                    generate_presigned_download_url as _gen_url,
                )
                logo_url = _gen_url(user_logo_key)
                lr = requests.get(logo_url, timeout=10)
                if lr.ok:
                    logo_ext = (
                        os.path.splitext(user_logo_key)[1].lower() or ".png"
                    )
                    user_logo_file = tmp_dir / f"logo{logo_ext}"
                    with open(user_logo_file, "wb") as f:
                        f.write(lr.content)
                    logo_path = str(user_logo_file)
            except Exception as e:
                logger.warning("Couldn't fetch user logo: %s", e)
        if not logo_path:
            try:
                from app.config import settings as _s
                cand = getattr(_s, "COMPANY_LOGO_PATH", None)
                if cand and os.path.isfile(cand):
                    logo_path = cand
            except Exception:
                logo_path = None

        translator_email = getattr(project, "_export_user_email", "") or ""

        project_meta = {
            "title": (
                getattr(project, "file_name", None)
                or "Certified Translation"
            ),
            "file_name": getattr(project, "file_name", None) or "",
            "source_language": getattr(project, "source_language", "") or "",
            "target_language": getattr(project, "target_language", "") or "",
            "translator_email": translator_email,
            "certification_date": datetime.utcnow().strftime(
                "%Y-%m-%d %H:%M UTC"
            ),
        }

        pdf_bytes = render_structured_export(
            source_kind=kind,
            original_path=str(src_path),
            pairs=pairs,
            project_meta=project_meta,
            company_logo_path=logo_path,
        )

        out = BytesIO()
        out.write(pdf_bytes)
        out.seek(0)
        return out
    except Exception as e:
        logger.warning(
            "Structured export rendering failed: %s", e, exc_info=True
        )
        return None
    finally:
        try:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass


def _resolve_cert_template(project, tmp_dir):
    """Return the raw bytes of the cert template DOCX the project is
    bound to, with {{tokens}} already substituted. Returns None when
    the project has no template, or when fetching/substituting fails
    (callers fall back to the hardcoded cert block in that case)."""
    template_id = getattr(project, "certification_template_id", None)
    if not template_id:
        return None

    try:
        from app.database import SessionLocal
        from app.models.certification import Certification
        from app.services.cert_template_service import (
            substitute_in_docx,
            build_substitution_values,
        )

        db = SessionLocal()
        try:
            cert = (
                db.query(Certification)
                .filter(Certification.id == template_id)
                .first()
            )
            if not cert:
                logger.info(
                    "Cert template %s not found — falling back to "
                    "hardcoded block",
                    template_id,
                )
                return None
            if not (cert.file_name or "").lower().endswith(".docx"):
                logger.info(
                    "Cert template %s isn't a DOCX (%s) — skipping "
                    "substitution",
                    template_id,
                    cert.file_name,
                )
                return None
            # File could live on local disk (legacy) or in Supabase
            # Storage (new). Try local first, then fall back to a
            # signed URL fetch from object storage.
            import os as _os
            template_bytes: bytes = b""
            if (
                cert.file_path
                and _os.path.exists(cert.file_path)
                and _os.path.isfile(cert.file_path)
            ):
                with open(cert.file_path, "rb") as f:
                    template_bytes = f.read()
            else:
                try:
                    from app.services.s3_service import (
                        generate_presigned_download_url,
                    )
                    import requests as _req
                    url = generate_presigned_download_url(cert.file_path)
                    r = _req.get(url, timeout=15)
                    r.raise_for_status()
                    template_bytes = r.content
                except Exception as e:
                    logger.warning(
                        "Couldn't fetch cert template from storage: %s",
                        e,
                    )
                    return None
            if not template_bytes:
                return None

            # Resolve the team for {{team_name}} / {{company_address}}.
            from app.models.team import Team
            team = (
                db.query(Team)
                .filter(Team.id == project.team_id)
                .first()
            )
            # The export pipeline stashes the user email + logo on
            # the project object as private attrs. Build a fake
            # "user" with the fields the substitution wants.
            class _U:
                pass
            fake_user = _U()
            fake_user.email = getattr(project, "_export_user_email", "") or ""
            fake_user.full_name = ""  # the email username is used as
                                      # fallback by build_substitution_values

            values = build_substitution_values(
                user=fake_user,
                project=project,
                team=team,
            )
            return substitute_in_docx(template_bytes, values)
        finally:
            db.close()
    except Exception as e:
        logger.warning(
            "Cert template substitution failed (using fallback): %s", e
        )
        return None


def _build_layout_docx_live(segments, project):
    """Build the export DOCX in the new STRUCTURED format — same shape
    as `_build_layout_pdf_live` but emitting .docx instead of .pdf.

        Page 1+      : original document embedded as image(s)
        Pages N+     : clean structured translation (alignment, bold,
                       italic, font family, placeholders all preserved)
        Final page   : certification block (optional company logo)
    """
    if project is None or not segments:
        return None
    kind = (getattr(project, "source_kind", "") or "").upper()
    if kind not in ("PDF", "IMAGE"):
        return None

    file_key = getattr(project, "file_path", None)
    if not file_key:
        return None

    import os
    import tempfile
    from datetime import datetime
    from pathlib import Path
    from app.services.s3_service import generate_presigned_download_url
    from app.services.structured_translation_renderer import (
        render_structured_docx_export,
        render_planned_docx_export,
    )
    import requests

    try:
        url = generate_presigned_download_url(file_key)
        r = requests.get(url, timeout=20)
        r.raise_for_status()
    except Exception as e:
        logger.warning("Couldn't fetch source for DOCX export: %s", e)
        return None

    tmp_dir = Path(tempfile.mkdtemp(prefix="traqexport_docx_"))
    try:
        ext = (
            ".pdf"
            if kind == "PDF"
            else os.path.splitext(file_key)[1] or ".jpg"
        )
        src_path = tmp_dir / f"source{ext}"
        with open(src_path, "wb") as f:
            f.write(r.content)

        pairs = []
        for s in segments:
            if not s.translated_text or not s.translated_text.strip():
                continue
            meta = dict(s.layout_meta or {})
            meta["source_text"] = s.source_text
            pairs.append((s.translated_text, meta))

        if not pairs:
            return None

        # Per-user logo first, then optional global fallback.
        logo_path = None
        user_logo_key = getattr(project, "_export_user_logo_key", None)
        if user_logo_key:
            try:
                logo_url = generate_presigned_download_url(user_logo_key)
                lr = requests.get(logo_url, timeout=10)
                if lr.ok:
                    logo_ext = (
                        os.path.splitext(user_logo_key)[1].lower() or ".png"
                    )
                    user_logo_file = tmp_dir / f"logo{logo_ext}"
                    with open(user_logo_file, "wb") as f:
                        f.write(lr.content)
                    logo_path = str(user_logo_file)
            except Exception as e:
                logger.warning("Couldn't fetch user logo: %s", e)
        if not logo_path:
            try:
                from app.config import settings as _s
                cand = getattr(_s, "COMPANY_LOGO_PATH", None)
                if cand and os.path.isfile(cand):
                    logo_path = cand
            except Exception:
                logo_path = None

        project_meta = {
            "title": (
                getattr(project, "file_name", None)
                or "Certified Translation"
            ),
            "file_name": getattr(project, "file_name", None) or "",
            "source_language": getattr(project, "source_language", "") or "",
            "target_language": getattr(project, "target_language", "") or "",
            "translator_email": (
                getattr(project, "_export_user_email", "") or ""
            ),
            "certification_date": datetime.utcnow().strftime(
                "%Y-%m-%d %H:%M UTC"
            ),
        }

        # If the project has a cert template selected, fetch + substitute
        # placeholders so the rebuild renderer can append it instead of
        # the hardcoded "I hereby certify..." block.
        cert_template_bytes = _resolve_cert_template(project, tmp_dir)
        if cert_template_bytes:
            project_meta["certification_template_bytes"] = cert_template_bytes

        # Layout-aware DOCX export is the DEFAULT. The planner asks
        # Claude to plan a DOCX skeleton mirroring the original
        # document's columns / tables / centered titles. On any
        # exception or empty result we fall through to the linear
        # renderer so exports never break. To force the linear path
        # set USE_LAYOUT_PLANNER=0 in the Railway env vars.
        docx_bytes = None
        if os.getenv("USE_LAYOUT_PLANNER", "1") != "0":
            try:
                docx_bytes = render_planned_docx_export(
                    source_kind=kind,
                    original_path=str(src_path),
                    pairs=pairs,
                    project_meta=project_meta,
                    company_logo_path=logo_path,
                )
            except Exception as e:
                logger.warning(
                    "Layout-aware DOCX render failed, falling back to linear: %s",
                    e,
                )
                docx_bytes = None
        if not docx_bytes:
            docx_bytes = render_structured_docx_export(
                source_kind=kind,
                original_path=str(src_path),
                pairs=pairs,
                project_meta=project_meta,
                company_logo_path=logo_path,
            )

        out = BytesIO()
        out.write(docx_bytes)
        out.seek(0)
        return out
    except Exception as e:
        logger.warning(
            "Structured DOCX export failed: %s", e, exc_info=True
        )
        return None
    finally:
        try:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass


def generate_docx(segments, user_email, project=None, user=None):
    if project is not None:
        try:
            project._export_user_email = user_email
            project._export_user_logo_key = (
                getattr(user, "logo_s3_key", None) if user else None
            )
        except Exception:
            pass

        # Prefer the new structured DOCX (mirrors the PDF export
        # structure: original on page 1, structured translation on
        # pages 2+, cert + logo at the end).
        structured = _build_layout_docx_live(segments, project)
        if structured is not None:
            return structured

        # Legacy fallback: the worker-written rebuilt DOCX prepended
        # to the cert block. Only triggers for DOCX source projects
        # (where the structured renderer can't run from an image-
        # only source).
        layout_doc = _try_layout_docx_from_project(project)
        if layout_doc is not None:
            return layout_doc

    doc = Document()
    # Translated body first…
    for seg in segments:
        if seg.translated_text:
            doc.add_paragraph(seg.translated_text)
    # …then the certification block at the end.
    doc.add_paragraph("")
    for line in build_certification(user_email):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    """Convert DOCX bytes to PDF using LibreOffice headless.

    Returns PDF bytes on success, None if LibreOffice isn't available
    or conversion fails (so the caller can fall back to the ReportLab
    PDF builder).

    This is how Export PDF stays visually identical to Export DOCX —
    we build ONE document (via the layout planner) and let LibreOffice
    render that exact DOCX into a PDF instead of running a separate
    ReportLab pipeline that drifts from the DOCX output.
    """
    import shutil
    import subprocess
    import tempfile
    from pathlib import Path

    # Locate LibreOffice. Different distros use different binary
    # names; we try the common ones in order.
    candidates = ["soffice", "libreoffice"]
    libreoffice_bin = next(
        (shutil.which(c) for c in candidates if shutil.which(c)), None
    )
    if not libreoffice_bin:
        logger.warning(
            "LibreOffice not found on PATH — PDF export will fall back "
            "to ReportLab. Install libreoffice on the deployment host "
            "to get DOCX-identical PDF output."
        )
        return None

    with tempfile.TemporaryDirectory(prefix="docx2pdf_") as tmp:
        tmp_dir = Path(tmp)
        docx_path = tmp_dir / "input.docx"
        docx_path.write_bytes(docx_bytes)

        try:
            # `--headless` runs without a UI. We pass an isolated
            # user-profile dir so concurrent calls don't collide on
            # the global LibreOffice profile (which would block).
            user_profile = tmp_dir / "lo_profile"
            user_profile.mkdir()
            env_arg = f"-env:UserInstallation=file://{user_profile}"

            proc = subprocess.run(
                [
                    libreoffice_bin,
                    env_arg,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_dir),
                    str(docx_path),
                ],
                capture_output=True,
                timeout=120,
                text=True,
            )
        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice conversion timed out after 120s")
            return None
        except Exception as e:
            logger.warning("LibreOffice invocation failed: %s", e)
            return None

        if proc.returncode != 0:
            logger.warning(
                "LibreOffice exited %d. stderr=%s",
                proc.returncode,
                (proc.stderr or "")[:500],
            )
            return None

        pdf_path = tmp_dir / "input.pdf"
        if not pdf_path.exists():
            logger.warning(
                "LibreOffice ran cleanly but produced no PDF at %s. "
                "stdout=%s",
                pdf_path,
                (proc.stdout or "")[:500],
            )
            return None

        logger.info("LibreOffice DOCX→PDF conversion succeeded")
        return pdf_path.read_bytes()


def generate_pdf(segments, user_email, project=None, user=None):
    """PDF export.

    Primary path: build the DOCX (which uses the layout planner with
    borders, tables, two-column rows etc.) and convert it to PDF via
    LibreOffice. This guarantees the PDF looks identical to the DOCX
    — they're literally the same document, one is rendered through
    the Office file format and one through PDF.

    Fallback path: if LibreOffice isn't installed (local dev), build
    the PDF with ReportLab via the existing _build_layout_pdf_live
    helper, and finally fall back to a flat paragraph stream.
    """
    # 1) Try the DOCX→LibreOffice→PDF path.
    if project is not None:
        try:
            docx_buf = generate_docx(
                segments, user_email, project=project, user=user
            )
        except Exception as e:
            logger.warning(
                "DOCX build failed inside PDF path, falling back: %s", e
            )
            docx_buf = None

        if docx_buf is not None:
            try:
                docx_bytes = docx_buf.getvalue()
            except AttributeError:
                # Already raw bytes
                docx_bytes = docx_buf

            pdf_bytes = _convert_docx_to_pdf(docx_bytes)
            if pdf_bytes:
                buf = BytesIO(pdf_bytes)
                buf.seek(0)
                return buf

    # 2) Fallback: ReportLab path mirroring the layout plan when we
    #    can (still tries the structured renderer), otherwise a flat
    #    paragraph stream.
    if project is not None:
        try:
            project._export_user_email = user_email
            project._export_user_logo_key = (
                getattr(user, "logo_s3_key", None) if user else None
            )
        except Exception:
            pass
        layout_pdf = _build_layout_pdf_live(segments, project)
        if layout_pdf is not None:
            return layout_pdf

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    content = []
    for seg in segments:
        if seg.translated_text:
            content.append(Paragraph(seg.translated_text, styles["Normal"]))
            content.append(Spacer(1, 10))
    content.append(Spacer(1, 20))
    for line in build_certification(user_email):
        content.append(Paragraph(line, styles["Normal"]))
        content.append(Spacer(1, 10))
    doc.build(content)
    buffer.seek(0)
    return buffer
