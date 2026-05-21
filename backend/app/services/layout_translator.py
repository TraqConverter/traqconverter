"""Layout-preserving translation pipeline.

Why this exists
---------------
The previous pipeline collapsed all whitespace into single spaces, regex-split
the result into "sentences", and emitted a flat stack of paragraphs. The
"output_file" uploaded to S3 was literally a copy of the original, untouched.

This module replaces both halves:

1. `extract_segments(...)` returns segments **with layout metadata** so that:
   - PDF blocks keep their bbox/font/page so we can put the translation
     back at the same coordinates.
   - DOCX paragraphs/table cells keep their position in the document tree
     so we can rewrite each in place — preserving boxes, tables, headings.
   - Image documents (passports/IDs/scans) are OCR'd at word level via
     pytesseract `image_to_data` and grouped into lines with bounding
     boxes so the translation can be drawn on a copy of the image.

2. `rebuild_output(...)` consumes the translated segments + layout metadata
   and writes a NEW file in the same format that closely resembles the
   source — table cells stay table cells, ID-style boxes stay boxed, etc.

Both halves degrade gracefully when extraction can't pull layout info
(scanned PDFs without OCR data, exotic encodings, etc.) — they fall back
to the simple paragraph stack used by the export router.
"""
from __future__ import annotations

import io
import logging
import os
from dataclasses import dataclass, field
from typing import Any

logger = logging.getLogger(__name__)


# ============================================================
# Public segment shape used between extractor and rebuilder.
# ============================================================

@dataclass
class ExtractedSegment:
    text: str
    layout: dict[str, Any] = field(default_factory=dict)


# ============================================================
# Font selection for non-Latin scripts.
# ----------------------------------------------------------------
# PyMuPDF's default "helv" font has no glyphs for CJK / Arabic /
# Cyrillic / Greek scripts — drawing those characters with it
# produces missing glyphs or a flat-out failure. PyMuPDF ships with
# CJK reserved font names that work without any file embedding;
# for other scripts we fall back to "helv" and accept that some
# characters may be missing until we wire a Noto font in.
# ============================================================
_CJK_FONTS = {
    "ja": "japan",       # Japanese sans-serif
    "jp": "japan",
    "ko": "korea",       # Korean sans-serif
    "kr": "korea",
    "zh": "china-s",     # Chinese (Simplified) sans-serif
}


def _pdf_font_for_language(target_lang: str | None) -> str:
    """Return a PyMuPDF font name suitable for the target language's
    primary script."""
    if not target_lang:
        return "helv"
    code = target_lang.lower()
    # zh-TW / zh-HK should use traditional Chinese
    if code.startswith("zh"):
        if "tw" in code or "hk" in code:
            return "china-t"
        return "china-s"
    base = code.split("-")[0]
    return _CJK_FONTS.get(base, "helv")


# ============================================================
# RTL (right-to-left) language support
# ----------------------------------------------------------------
# Arabic and Hebrew (plus Farsi, Urdu, Yiddish) read right-to-left.
# Three things need to happen so PyMuPDF renders them correctly:
#
#   1. Arabic letters have positional forms (initial, medial, final,
#      isolated) — arabic_reshaper converts logical characters into
#      visually-joined glyphs. Hebrew doesn't need this step.
#   2. Bidirectional algorithm (UAX #9) reorders the text for visual
#      display so RTL runs appear right-to-left while embedded LTR runs
#      (numbers, English words) keep their natural order. python-bidi
#      implements this.
#   3. The text-box must be right-aligned (PyMuPDF align=2) so the line
#      starts at the right edge of the box, where RTL readers expect it.
#
# Both libraries are optional — if they're not installed we still apply
# right-alignment, but Arabic glyphs won't join and bidi runs won't
# reorder correctly.
# ============================================================
_RTL_LANGS = {"ar", "he", "fa", "ur", "yi", "ji"}


def _is_rtl_language(target_lang: str | None) -> bool:
    if not target_lang:
        return False
    base = target_lang.lower().split("-")[0]
    return base in _RTL_LANGS


def _shape_rtl(text: str, target_lang: str | None) -> str:
    """Apply Arabic reshaping (if applicable) and the Unicode
    bidirectional algorithm so the text displays correctly in a
    left-to-right rendering engine like PyMuPDF."""
    if not _is_rtl_language(target_lang) or not text:
        return text

    base = (target_lang or "").lower().split("-")[0]
    shaped = text
    try:
        if base in ("ar", "fa", "ur"):
            import arabic_reshaper
            shaped = arabic_reshaper.reshape(shaped)
    except Exception:
        # arabic-reshaper missing — fall through with original text.
        pass

    try:
        from bidi.algorithm import get_display
        shaped = get_display(shaped)
    except Exception:
        # python-bidi missing — text will read in logical order.
        pass

    return shaped


# Candidate paths to search for a font that supports the target
# language's script. PyMuPDF's reserved CJK names work without files
# but Arabic/Hebrew/Hindi need an actual TTF. We probe common system
# locations and register the first one we find.
_SCRIPT_FONT_CANDIDATES: dict[str, list[str]] = {
    "ar": [
        "C:/Windows/Fonts/arabtype.ttf",
        "C:/Windows/Fonts/arial.ttf",  # has basic Arabic glyphs
        "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
    ],
    "fa": [  # Persian uses Arabic script
        "C:/Windows/Fonts/arabtype.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
    ],
    "ur": [
        "C:/Windows/Fonts/arabtype.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
    ],
    "he": [
        "C:/Windows/Fonts/david.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansHebrew-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ],
    "hi": [
        "C:/Windows/Fonts/mangal.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
    ],
    "th": [
        "C:/Windows/Fonts/tahoma.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf",
    ],
}


def _find_script_font_file(target_lang: str | None) -> str | None:
    if not target_lang:
        return None
    base = target_lang.lower().split("-")[0]
    for path in _SCRIPT_FONT_CANDIDATES.get(base, []):
        if os.path.isfile(path):
            return path
    return None


def _register_script_font(page, target_lang: str | None) -> str | None:
    """If the target language uses a non-Latin script that's not one of
    PyMuPDF's built-in CJK fonts, register a system TTF and return the
    name we registered it under. Returns None if no suitable file was
    found (caller will fall back to helv)."""
    path = _find_script_font_file(target_lang)
    if not path:
        return None
    name = f"script-{(target_lang or '').lower().split('-')[0]}"
    try:
        page.insert_font(fontname=name, fontfile=path)
        return name
    except Exception:
        return None


# ============================================================
# Source kind detection
# ============================================================

def detect_source_kind(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return "PDF"
    if ext == ".docx":
        return "DOCX"
    if ext in (".jpg", ".jpeg", ".png"):
        return "IMAGE"
    if ext == ".txt":
        return "TXT"
    return "UNKNOWN"


# ============================================================
# EXTRACTORS
# ============================================================

def _is_scanned_pdf(doc) -> bool:
    """Heuristic to detect 'scanned image PDFs' (CamScanner, photocopier
    output, etc.) where the embedded text layer is noisy OCR.

    A page is considered scanned when:
      * it contains at least one image block
      * AND the image area is >= 50% of the page area

    If at least one page in the document is a scan, we prefer fresh OCR.
    """
    try:
        for page in doc:
            page_area = max(1.0, page.rect.width * page.rect.height)
            data = page.get_text("dict")
            image_area = 0.0
            for block in data.get("blocks", []):
                # type 1 == image block in PyMuPDF
                if block.get("type") == 1:
                    bbox = block.get("bbox") or [0, 0, 0, 0]
                    w = max(0.0, bbox[2] - bbox[0])
                    h = max(0.0, bbox[3] - bbox[1])
                    image_area += w * h
            if image_area / page_area >= 0.5:
                return True
    except Exception:
        return False
    return False


def _sample_background_color(
    page,
    rect,
    pix=None,
    margin: float = 3.0,
) -> tuple[float, float, float]:
    """Sample pixel colors just outside the bbox edges and return the
    average as an (r, g, b) tuple in 0..1 range. The caller can pass a
    pre-rendered pixmap of the whole page (`pix`) so we don't re-render
    once per block.

    Falls back to white (1, 1, 1) when sampling can't run (no pixmap,
    bbox out of bounds, etc.).
    """
    try:
        if pix is None:
            pix = page.get_pixmap(dpi=72)  # 72dpi → 1pt = 1px
        n_components = pix.n  # 3 (RGB) or 4 (RGBA)
        page_w_px = pix.width
        page_h_px = pix.height

        # Build sample points just OUTSIDE the bbox on each side.
        sample_pts = []
        # Edge points along top and bottom.
        for fx in (0.1, 0.3, 0.5, 0.7, 0.9):
            x = int(rect.x0 + fx * (rect.x1 - rect.x0))
            sample_pts.append((x, int(rect.y0 - margin)))
            sample_pts.append((x, int(rect.y1 + margin)))
        # Edge points along left and right.
        for fy in (0.3, 0.5, 0.7):
            y = int(rect.y0 + fy * (rect.y1 - rect.y0))
            sample_pts.append((int(rect.x0 - margin), y))
            sample_pts.append((int(rect.x1 + margin), y))

        rs, gs, bs = [], [], []
        samples = pix.samples
        for x, y in sample_pts:
            if 0 <= x < page_w_px and 0 <= y < page_h_px:
                idx = (y * page_w_px + x) * n_components
                rs.append(samples[idx])
                gs.append(samples[idx + 1])
                bs.append(samples[idx + 2])

        if not rs:
            return (1.0, 1.0, 1.0)

        # Median is more robust than mean against e.g. nearby ink pixels.
        rs.sort(); gs.sort(); bs.sort()
        m = len(rs) // 2
        r = rs[m] / 255.0
        g = gs[m] / 255.0
        b = bs[m] / 255.0
        return (r, g, b)
    except Exception:
        return (1.0, 1.0, 1.0)


def _bbox_overlap_ratio(a: list[float], b: list[float]) -> float:
    """Returns the fraction of the smaller bbox that overlaps the larger.
    0.0 means no overlap, 1.0 means one bbox is fully inside the other."""
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0 = max(ax0, bx0)
    iy0 = max(ay0, by0)
    ix1 = min(ax1, bx1)
    iy1 = min(ay1, by1)
    iw = max(0.0, ix1 - ix0)
    ih = max(0.0, iy1 - iy0)
    inter = iw * ih
    if inter <= 0:
        return 0.0
    a_area = max(1.0, (ax1 - ax0) * (ay1 - ay0))
    b_area = max(1.0, (bx1 - bx0) * (by1 - by0))
    return inter / min(a_area, b_area)


def _on_same_line(a: list[float], b: list[float]) -> bool:
    """Return True if two bboxes are on the same text line.

    Uses the vertical centres: they must be within ~30% of the average
    bbox height of each other. Two vertically-stacked lines have centres
    separated by roughly one full line-height, so this filter keeps them
    apart even when their bboxes touch or slightly overlap.
    """
    a_h = max(1.0, a[3] - a[1])
    b_h = max(1.0, b[3] - b[1])
    a_cy = (a[1] + a[3]) / 2
    b_cy = (b[1] + b[3]) / 2
    avg_h = (a_h + b_h) / 2
    return abs(a_cy - b_cy) < avg_h * 0.3


def _merge_overlapping_segments(
    segments: list[ExtractedSegment],
) -> list[ExtractedSegment]:
    """Merge segments whose bboxes overlap on the SAME text line.

    The classic problem is two blocks that share the same row (a label
    and its value side by side, or PyMuPDF returning two spans for one
    line). Those should be merged because painting two translations on
    top of each other looks awful.

    The classic anti-case is two consecutive lines that touch slightly
    vertically — those must NOT be merged because the resulting tall
    bbox inflates the font size and creates big visible gaps. The
    "same line" check below uses vertical-centre distance to keep
    stacked lines independent.
    """
    if not segments:
        return segments

    by_page: dict[int, list[ExtractedSegment]] = {}
    for s in segments:
        page = int((s.layout or {}).get("page", 0))
        by_page.setdefault(page, []).append(s)

    merged_all: list[ExtractedSegment] = []
    for page, page_segs in by_page.items():
        bins: list[ExtractedSegment] = []
        for s in page_segs:
            bbox = (s.layout or {}).get("bbox")
            if not bbox:
                bins.append(s)
                continue
            placed = False
            for b in bins:
                bb = (b.layout or {}).get("bbox")
                if not bb:
                    continue
                if (
                    _on_same_line(bbox, bb)
                    and _bbox_overlap_ratio(bbox, bb) >= 0.3
                ):
                    nb = [
                        min(bb[0], bbox[0]),
                        min(bb[1], bbox[1]),
                        max(bb[2], bbox[2]),
                        max(bb[3], bbox[3]),
                    ]
                    b.layout["bbox"] = nb
                    b.text = (b.text + " " + s.text).strip()
                    placed = True
                    break
            if not placed:
                bins.append(s)
        merged_all.extend(bins)

    return merged_all


def _map_pdf_font(font_name: str | None, flags: int) -> str:
    """Map an arbitrary PDF font name + style flags to one of PyMuPDF's
    14 built-in font names (so glyphs render without needing to embed
    a TTF).

    PyMuPDF span flag bits:
        bit 0 (1)   superscript
        bit 1 (2)   italic
        bit 2 (4)   serif
        bit 3 (8)   monospaced
        bit 4 (16)  bold
    """
    is_bold = bool(flags & 16)
    is_italic = bool(flags & 2)
    is_serif = bool(flags & 4)
    is_mono = bool(flags & 8)

    fn = (font_name or "").lower()

    # Monospaced → Courier family
    if is_mono or any(k in fn for k in ("courier", "mono", "consolas")):
        if is_bold and is_italic:
            return "cobi"
        if is_bold:
            return "cobo"
        if is_italic:
            return "coit"
        return "cour"

    # Serif → Times family
    serif_hits = ("times", "garamond", "georgia", "serif", "roman", "minion", "cambria")
    if is_serif or any(k in fn for k in serif_hits):
        if is_bold and is_italic:
            return "tibi"
        if is_bold:
            return "tibo"
        if is_italic:
            return "tiit"
        return "tiro"

    # Default → Helvetica family
    if is_bold and is_italic:
        return "hebi"
    if is_bold:
        return "hebo"
    if is_italic:
        return "heit"
    return "helv"


def _extract_pdf(file_path: str) -> list[ExtractedSegment]:
    """Extract source text at LINE level (not block level).

    Why per-line
    ------------
    PyMuPDF text blocks group several visually-consecutive lines into one
    rectangle. Translating block by block and redrawing the whole block
    inside its bounding rect loses the original line spacing, alignment,
    and breaks. Per-line extraction preserves:
      * the exact bbox of every physical line (so we redact and redraw
        in-place — no drift)
      * the line's dominant font, size, colour, bold/italic flags (so
        the translation visually matches)

    Scanned PDFs (CamScanner-style: image-only pages plus a noisy OCR
    text layer) are detected and routed through fresh pytesseract OCR
    so we never trust an unreliable embedded OCR.
    """
    import fitz

    doc = fitz.open(file_path)
    try:
        if _is_scanned_pdf(doc):
            logger.info(
                "PDF appears to be a scan — using fresh OCR instead of "
                "the embedded text layer."
            )
            doc.close()
            return _extract_pdf_via_ocr(file_path)

        out: list[ExtractedSegment] = []
        for page_index, page in enumerate(doc):
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue  # skip image blocks

                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue

                    # Concatenate span text in order; pick the
                    # dominant font properties from the span that
                    # covers the most characters (so a 1-char span
                    # in a different font doesn't override the line).
                    line_text = ""
                    dominant_span = max(
                        spans,
                        key=lambda s: len(s.get("text", "") or ""),
                    )
                    for s in spans:
                        line_text += s.get("text", "") or ""
                    line_text = line_text.strip()
                    if not line_text:
                        continue

                    font_size = float(dominant_span.get("size") or 11.0)
                    color = int(dominant_span.get("color") or 0)
                    font_name = dominant_span.get("font") or ""
                    flags = int(dominant_span.get("flags") or 0)
                    ascender = float(dominant_span.get("ascender") or 0.8)
                    descender = float(dominant_span.get("descender") or -0.2)

                    bbox = list(line.get("bbox", [0, 0, 0, 0]))
                    out.append(
                        ExtractedSegment(
                            text=line_text,
                            layout={
                                "kind": "pdf_line",
                                "page": page_index,
                                "bbox": bbox,
                                "font_size": font_size,
                                "color": color,
                                "font_name": font_name,
                                "flags": flags,
                                "ascender": ascender,
                                "descender": descender,
                            },
                        )
                    )
    finally:
        try:
            doc.close()
        except Exception:
            pass

    if not out:
        out = _extract_pdf_via_ocr(file_path)
    return out

    # Fallback OCR if nothing extracted at all.
    if not out:
        out = _extract_pdf_via_ocr(file_path)

    return out


def _extract_pdf_via_ocr(file_path: str) -> list[ExtractedSegment]:
    """Render each page and OCR with bounding boxes.

    Renders at 2x scale so small print (phone numbers, registration codes,
    footer disclaimers) is legible to tesseract — the default 1x render
    drops most fine print, leaving the translation incomplete.
    """
    import fitz
    import pytesseract
    from PIL import Image

    out: list[ExtractedSegment] = []
    doc = fitz.open(file_path)
    try:
        # 2x render so OCR can read small print confidently. The bboxes
        # come back in the *rendered* coordinate space; we convert to the
        # PDF page's native coordinates so the rebuild draws translations
        # at the correct positions.
        OCR_SCALE = 2.0
        matrix = fitz.Matrix(OCR_SCALE, OCR_SCALE)
        for page_index, page in enumerate(doc):
            pix = page.get_pixmap(matrix=matrix)
            img = Image.frombytes(
                "RGB", [pix.width, pix.height], pix.samples
            )
            try:
                data = pytesseract.image_to_data(
                    img, output_type=pytesseract.Output.DICT
                )
            except Exception as e:
                logger.warning(f"OCR failed on page {page_index}: {e}")
                continue

            # Group by (block_num, par_num, line_num) → one segment per line
            lines: dict[tuple, dict] = {}
            n = len(data.get("text", []))
            for i in range(n):
                txt = (data["text"][i] or "").strip()
                if not txt:
                    continue
                key = (
                    data["block_num"][i],
                    data["par_num"][i],
                    data["line_num"][i],
                )
                entry = lines.setdefault(
                    key,
                    {
                        "words": [],
                        "x": int(data["left"][i]),
                        "y": int(data["top"][i]),
                        "right": int(data["left"][i] + data["width"][i]),
                        "bottom": int(data["top"][i] + data["height"][i]),
                    },
                )
                entry["words"].append(txt)
                entry["x"] = min(entry["x"], int(data["left"][i]))
                entry["y"] = min(entry["y"], int(data["top"][i]))
                entry["right"] = max(
                    entry["right"], int(data["left"][i] + data["width"][i])
                )
                entry["bottom"] = max(
                    entry["bottom"], int(data["top"][i] + data["height"][i])
                )

            for line_idx, info in enumerate(lines.values()):
                line_text = " ".join(info["words"]).strip()
                if not line_text:
                    continue
                # Convert OCR bbox from rendered coords back to PDF
                # native coords (so the rebuild draws translations at
                # the right positions on the same-sized page).
                inv = 1.0 / OCR_SCALE
                # Derive a reasonable font size from bbox height. Cap at
                # 14pt so an unusually tall bbox (caused by ascenders/
                # descenders, stamps, or noisy OCR grouping) doesn't
                # render the translation at giant body-paragraph height.
                derived = (info["bottom"] - info["y"]) * inv * 0.6
                font_size = max(8.0, min(14.0, derived))
                out.append(
                    ExtractedSegment(
                        text=line_text,
                        layout={
                            "kind": "pdf_ocr_line",
                            "page": page_index,
                            "bbox": [
                                info["x"] * inv,
                                info["y"] * inv,
                                info["right"] * inv,
                                info["bottom"] * inv,
                            ],
                            "line": line_idx,
                            "font_size": font_size,
                        },
                    )
                )
    finally:
        doc.close()

    return out


def _extract_docx(file_path: str) -> list[ExtractedSegment]:
    """Walk paragraphs (top-level) and table cells in document order."""
    from docx import Document

    out: list[ExtractedSegment] = []
    doc = Document(file_path)

    # Top-level paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        out.append(
            ExtractedSegment(
                text=text,
                layout={
                    "kind": "docx_paragraph",
                    "para_index": para_idx,
                },
            )
        )

    # Tables — translate every cell, every paragraph inside the cell
    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for cp_idx, cp in enumerate(cell.paragraphs):
                    text = cp.text.strip()
                    if not text:
                        continue
                    out.append(
                        ExtractedSegment(
                            text=text,
                            layout={
                                "kind": "docx_table_cell",
                                "table_index": tbl_idx,
                                "row": row_idx,
                                "col": col_idx,
                                "para_in_cell": cp_idx,
                            },
                        )
                    )

    return out


def _looks_like_mrz(text: str) -> bool:
    """Detect Machine-Readable Zone lines (passport / ID card bottom).
    These are encoded data, not real language, and translating them
    produces nonsense. Pattern: long line of caps + digits + '<' chars."""
    s = (text or "").strip().replace(" ", "")
    if len(s) < 20:
        return False
    # MRZ lines use only A-Z, 0-9, and < ; have ≥2 chevron fillers.
    if s.count("<") >= 2 and all(c.isalnum() or c == "<" for c in s):
        return True
    return False


def _extract_image_via_claude(file_path: str) -> list[ExtractedSegment] | None:
    """Try Claude Vision OCR first. Returns None to signal 'fall back'."""
    try:
        from app.services.claude_vision_ocr import is_available, ocr_image
    except Exception:
        return None
    if not is_available():
        return None
    try:
        lines = ocr_image(file_path)
    except Exception as e:
        logger.warning("Claude Vision OCR raised: %s", e)
        return None
    if not lines:
        return None

    out: list[ExtractedSegment] = []
    for idx, line in enumerate(lines):
        text = (line.get("text") or "").strip()
        if not text:
            continue
        if _looks_like_mrz(text):
            continue
        bbox = line.get("bbox")
        if not bbox or len(bbox) != 4:
            continue
        x0, y0, x1, y1 = bbox
        h = max(1.0, y1 - y0)
        # Map Claude's relative "size" hint to an approximate point
        # size so the renderer keeps hierarchy (title vs body vs
        # footer fine print).
        rel_size = (line.get("size") or "normal").lower()
        size_map = {
            "small": 8.0,
            "normal": 10.5,
            "large": 13.0,
            "xlarge": 16.0,
        }
        font_size = size_map.get(rel_size, max(7.0, min(12.0, h * 0.7)))
        # Bold/italic via PyMuPDF span-flag bits so the existing
        # rendering paths read it the same way as PDF segments.
        flags = 0
        if line.get("bold"):
            flags |= 16
        if line.get("italic"):
            flags |= 2
        out.append(
            ExtractedSegment(
                text=text,
                layout={
                    "kind": "image_line",
                    "page": 0,
                    "bbox": [float(x0), float(y0), float(x1), float(y1)],
                    "line": idx,
                    "font_size": font_size,
                    "ocr_source": "claude",
                    "claude_kind": line.get("kind") or "text",
                    "placeholder_kind": line.get("placeholder_kind"),
                    "alignment": line.get("alignment") or "left",
                    "bold": bool(line.get("bold")),
                    "italic": bool(line.get("italic")),
                    "all_caps": bool(line.get("all_caps")),
                    "size_hint": rel_size,
                    "flags": flags,
                    # Claude's font-family classification used by the
                    # renderer to pick Times / Helvetica / Courier.
                    "font_family": line.get("font_family") or "",
                },
            )
        )
    return out if out else None


def _extract_image(file_path: str) -> list[ExtractedSegment]:
    """OCR an image into translatable line segments.

    Pipeline:
      1. Try Claude Vision OCR first (cleaner reads on ID/passport/
         certificate documents). Only runs when ANTHROPIC_API_KEY is set.
      2. Fall back to tesseract when Claude isn't available or returns
         nothing usable.

    The tesseract path is the same proven implementation as before:
      * 2400px upscale before OCR
      * per-word confidence filter (conf < 55 dropped)
      * MRZ pattern skip
      * median-word-height font sizing capped 7-12pt
    """
    claude_segs = _extract_image_via_claude(file_path)
    if claude_segs:
        return claude_segs

    import pytesseract
    from PIL import Image

    image = Image.open(file_path).convert("RGB")
    orig_w, orig_h = image.size

    # ----------------------------------------------------------------
    # 1. Upscale for better OCR. Tesseract's accuracy improves
    # substantially on inputs ≥ 2000px on the long edge.
    # ----------------------------------------------------------------
    long_edge = max(orig_w, orig_h)
    if long_edge < 2400:
        scale = 2400 / long_edge
        new_w = int(orig_w * scale)
        new_h = int(orig_h * scale)
        ocr_img = image.resize((new_w, new_h), Image.LANCZOS)
    else:
        scale = 1.0
        ocr_img = image

    try:
        data = pytesseract.image_to_data(
            ocr_img, output_type=pytesseract.Output.DICT
        )
    except Exception as e:
        raise Exception(f"OCR failed: {e}")

    inv_scale = 1.0 / scale  # rendered → source coords

    # ----------------------------------------------------------------
    # 2. Group words → lines, applying confidence filter.
    # ----------------------------------------------------------------
    lines: dict[tuple, dict] = {}
    n = len(data.get("text", []))
    MIN_CONF = 55  # tesseract returns -1 for "no confidence" entries
    for i in range(n):
        txt = (data["text"][i] or "").strip()
        if not txt:
            continue
        try:
            conf = int(data.get("conf", [-1])[i])
        except Exception:
            conf = -1
        if conf >= 0 and conf < MIN_CONF:
            continue  # low-confidence noise — drop word
        key = (
            data["block_num"][i],
            data["par_num"][i],
            data["line_num"][i],
        )
        x = int(data["left"][i])
        y = int(data["top"][i])
        w = int(data["width"][i])
        h = int(data["height"][i])
        entry = lines.setdefault(
            key,
            {
                "words": [],
                "x": x,
                "y": y,
                "right": x + w,
                "bottom": y + h,
                "heights": [],
            },
        )
        entry["words"].append(txt)
        entry["x"] = min(entry["x"], x)
        entry["y"] = min(entry["y"], y)
        entry["right"] = max(entry["right"], x + w)
        entry["bottom"] = max(entry["bottom"], y + h)
        entry["heights"].append(h)

    out: list[ExtractedSegment] = []
    for line_idx, info in enumerate(lines.values()):
        line_text = " ".join(info["words"]).strip()
        if not line_text:
            continue
        if _looks_like_mrz(line_text):
            continue  # MRZ codes are not translatable text

        # Convert from upscaled-image coords back to source-image coords.
        x0 = info["x"] * inv_scale
        y0 = info["y"] * inv_scale
        x1 = info["right"] * inv_scale
        y1 = info["bottom"] * inv_scale

        # Font size — use median word-height to stay robust against
        # one tall ascender skewing the bbox. Cap tightly so the
        # rebuild never paints absurdly-large translations.
        heights = sorted(info["heights"]) or [16]
        med_h = heights[len(heights) // 2] * inv_scale
        font_size = max(7.0, min(12.0, med_h * 0.7))

        out.append(
            ExtractedSegment(
                text=line_text,
                layout={
                    "kind": "image_line",
                    "page": 0,
                    "bbox": [x0, y0, x1, y1],
                    "line": line_idx,
                    "font_size": font_size,
                },
            )
        )
    return out


def _extract_txt(file_path: str) -> list[ExtractedSegment]:
    """Plain text — preserve paragraph splits (blank-line separated)."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    out: list[ExtractedSegment] = []
    for idx, para in enumerate(raw.split("\n\n")):
        text = para.strip()
        if not text:
            continue
        out.append(
            ExtractedSegment(
                text=text,
                layout={"kind": "txt_paragraph", "para_index": idx},
            )
        )
    return out


def extract_segments(file_path: str) -> tuple[str, list[ExtractedSegment]]:
    """Return (source_kind, segments). Segments preserve enough metadata
    that rebuild_output can reconstruct a similar-looking file."""
    kind = detect_source_kind(file_path)
    if kind == "PDF":
        segs = _extract_pdf(file_path)
        return kind, _merge_overlapping_segments(segs)
    if kind == "DOCX":
        return kind, _extract_docx(file_path)
    if kind == "IMAGE":
        segs = _extract_image(file_path)
        return kind, _merge_overlapping_segments(segs)
    if kind == "TXT":
        return kind, _extract_txt(file_path)
    raise Exception(f"Unsupported file type for {file_path}")


# ============================================================
# REBUILDERS — produce a file that resembles the source.
# ============================================================

import re as _re_noise

# Common short words across the languages we typically translate. If a block
# of supposed prose contains zero of these, it's almost certainly OCR noise,
# microprint, or scrambled security text — not real natural language.
_REAL_WORDS = {
    # English
    "the", "and", "for", "are", "with", "this", "that", "from", "have", "been",
    "will", "you", "your", "all", "but", "not", "any", "can", "has", "her",
    "him", "his", "may", "new", "now", "one", "out", "see", "two", "way",
    "who", "how", "its", "our", "use", "did", "day", "get", "name", "of",
    "is", "to", "be", "in", "on", "by", "as", "an", "or", "if", "do", "at",
    "we", "no", "so", "up", "us", "all", "ltd", "inc", "llc", "corp", "co",
    "registered", "training", "certificate", "certified", "issued", "expiry",
    "date", "competent", "competence", "issue", "this", "page", "name",
    "country", "passport", "type", "code", "given", "names", "surname",
    "nationality", "place", "birth", "sex", "authorities", "responsible",
    # Italian
    "che", "per", "non", "una", "del", "della", "con", "questo", "tutto",
    "essere", "sono", "stato", "dal", "delle", "dei", "alla", "anche", "come",
    "ancora", "ogni", "molto", "nella", "tutti", "loro", "ora", "oltre", "il",
    "la", "le", "lo", "gli", "uno", "una", "ed", "ma", "se", "si", "ho",
    # Spanish
    "que", "para", "este", "todo", "ser", "los", "las", "una", "muy", "mas",
    "como", "pero", "sus", "yo", "el", "ella", "nosotros", "su", "y", "es",
    "en", "un", "una", "del", "al", "se", "no", "se", "lo", "le", "me", "te",
    # French
    "pour", "dans", "ses", "des", "les", "elle", "ne", "pas", "ou", "ce",
    "qui", "sur", "avec", "par", "tout", "sans", "tres", "bien", "le", "la",
    "il", "et", "est", "un", "une", "au", "aux", "en", "de", "du", "ces",
    # German
    "der", "die", "das", "und", "ist", "ein", "eine", "nicht", "mit", "auf",
    "von", "im", "zu", "es", "sich", "auch", "wird", "kann", "uber", "muss",
    "den", "dem", "des", "wie", "wo", "wer", "was", "war", "hat", "ich",
    # Portuguese / Dutch / others — common short words
    "para", "como", "mas", "tem", "foi", "voor", "aan", "naar", "een", "het",
    "zijn", "deze", "deze", "voor", "ook", "bij",
}


# Common letter trigrams across the Latin-script languages we translate
# in/out of. Real natural-language words almost always contain at least
# one of these. Words that contain ZERO of them (and are >= 5 chars) are
# almost certainly OCR gibberish — "anauyafe", "eletrale", "statiton",
# "ceraacennen", "raeewanh" all fail this test.
_COMMON_TRIGRAMS = {
    # English
    "the", "ing", "ion", "ent", "ati", "for", "and", "her", "ate", "are",
    "ess", "est", "ist", "men", "tio", "ant", "all", "our", "his", "out",
    "ter", "ers", "res", "con", "pro", "com", "per", "ver", "tra", "ble",
    "ous", "ave", "ive", "ity", "ies", "ome", "ide", "ine", "one", "ore",
    # Italian
    "che", "ato", "are", "del", "ane", "ess", "ali", "ani", "ina", "ito",
    "ari", "ica", "tro", "ggi", "zio", "ttu", "ent", "ust", "età", "ela",
    "ele", "ide", "ola", "olo", "uti", "uto", "ure", "ura", "iva", "ivo",
    # Spanish
    "ado", "iar", "tiv", "ció", "ido", "ada", "ido", "ene", "ido",
    "ría", "ros", "tar", "tor", "ica", "iva",
    # French
    "ant", "tio", "ire", "ies", "eur", "eux", "ene", "ard", "ais",
    "tre", "ous", "que", "lle", "ble", "ité", "ais", "fra",
    # German
    "ich", "ein", "ung", "der", "die", "und", "sch", "ber", "lic",
    "ist", "auf", "vor", "über", "bei", "aus", "ren",
    # Portuguese / Dutch / generic
    "ção", "ade", "ent", "ada", "een", "het", "een", "voor",
}


# Bigrams that essentially never occur in any natural Latin-script
# language. Their presence in a 5+ char word almost certainly means OCR
# garbage. "uy" mid-word catches "anauyafe"; "vy" catches OCR'd headers
# like "vy, Stato. c"; "yj"/"jq" catch random keyboard noise.
_RARE_BIGRAMS = (
    "uy", "yj", "jq", "qx", "xz", "zx", "vp", "pv",
    "wj", "jw", "kx", "xk", "fz", "zf", "fq", "qf",
)


def _word_looks_natural(word: str) -> bool:
    """Returns True if a word looks like natural language.

    Short words (< 5 chars) are always considered natural since they
    don't have room for the test.

    A 5+ char word is considered NOT natural if it contains any of our
    rare-bigram markers (strong gibberish signal). Otherwise it must
    contain at least one common natural-language trigram.
    """
    w = word.lower()
    if len(w) < 5:
        return True
    for rb in _RARE_BIGRAMS:
        if rb in w:
            return False
    for tg in _COMMON_TRIGRAMS:
        if tg in w:
            return True
    return False


def _is_noise_text(text: str, font_size: float | None = None) -> bool:
    """Decide if a source-text block is noise we should skip drawing.

    Conservative on the obvious junk (microprint, MRZ filler, character
    runs, 5+ consonant streams), then a softer linguistic check: if the
    block contains 2+ longer alphabetic words and NONE of them look like
    real natural-language words (no common trigrams), it's gibberish.
    """
    s = (text or "").strip()
    if len(s) < 2:
        return True

    # Microprint / security text. Body text is almost always >= 8pt.
    if font_size is not None and float(font_size) < 7.0:
        return True

    # MRZ filler ("<<<<<<...")
    if s.count("<") >= max(3, len(s) // 4):
        return True

    # Mostly identical-character runs ("MMMMMMM", "----")
    unique = len(set(s.replace(" ", "")))
    if unique <= 2 and len(s) >= 4:
        return True

    raw_words = [w.strip(".,;:!?()[]{}\"'`|/\\-_@") for w in s.split()]

    # 5+ consecutive consonants inside any single word ⇒ gibberish.
    for w in raw_words:
        if not w or len(w) < 5 or not all(c.isalpha() for c in w):
            continue
        if _re_noise.search(r"[bcdfghjklmnpqrstvwxyz]{5,}", w.lower()):
            return True

    long_alpha = [
        w for w in raw_words
        if w and len(w) >= 5 and all(c.isalpha() for c in w)
    ]

    # ANY single word with rare bigrams (uy, yj, qz, etc.) poisons the
    # block — these patterns essentially never occur in real Latin-
    # script natural language and almost always come from stylised-text
    # OCR errors. One "anauyafe" is enough to flag the block.
    for w in long_alpha:
        wl = w.lower()
        for rb in _RARE_BIGRAMS:
            if rb in wl:
                return True

    # Soft linguistic check. If we have 2+ long words and not a single
    # one looks like real natural-language, treat as gibberish.
    if len(long_alpha) >= 2:
        if not any(_word_looks_natural(w) for w in long_alpha):
            return True

    return False


def _rebuild_pdf(
    original_path: str,
    output_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
    target_lang: str | None = None,
) -> None:
    """In-place overlay — open the original PDF and replace each source
    text block with its translation at the SAME font size, position, and
    color. Watermarks, logos, signatures, borders and the rest of the
    visual layout are preserved as-is.

    Font sizing
    -----------
    Translations are drawn at the *source* font size (no auto-shrinking
    to fit). If the translation is longer than the source bbox we expand
    the rectangle downward rather than miniaturising the text — keeps
    typography visually consistent with the original.

    Noise filter
    ------------
    Source blocks that look like OCR'd watermark/microprint/security
    gibberish are skipped so we don't paint mistranslated junk where the
    original showed decorative noise.
    """
    logger.info("[pdf-rebuild v6] line-level overlay with font matching")
    import fitz

    # Open the original — we modify a COPY of every page rather than building
    # a fresh canvas, so watermarks / images / borders / signatures stay.
    # Pick the font that supports the target language's script. CJK
    # languages (Japanese, Korean, Chinese) use PyMuPDF's reserved CJK
    # font names. Other non-Latin scripts (Arabic, Hebrew, Hindi, Thai)
    # are registered from a system TTF per-page below. Otherwise we
    # fall back to "helv" which covers Latin/Cyrillic/Greek.
    fontname = _pdf_font_for_language(target_lang)
    is_rtl = _is_rtl_language(target_lang)

    doc = fitz.open(original_path)
    try:
        # Register Arabic/Hebrew/Hindi/Thai system fonts once per page
        # (caching avoids re-registering on every block).
        registered_fonts: dict[int, str] = {}

        # Pre-render every page once so background sampling reads the
        # ORIGINAL pixels (before any of our redactions remove text).
        page_pixmaps = {}
        for idx, p in enumerate(doc):
            try:
                page_pixmaps[idx] = p.get_pixmap(dpi=72)
            except Exception:
                page_pixmaps[idx] = None

        # ============================================================
        # PASS 1 — REDACT the source text for every bbox we plan to
        # paint over. add_redact_annot physically removes the text from
        # the PDF stream (not just paints over it) so the output PDF
        # contains ONLY the translation, not Italian + English overlaid.
        #
        # CRITICAL: apply_redactions defaults to removing images and
        # vector graphics within the rectangle too. For form-heavy
        # documents (Italian tax forms, certificates) that would strip
        # the box outlines, dividers and field borders along with the
        # text — leaving an empty skeleton. We pass IMAGE_NONE +
        # LINE_ART_NONE so the redaction only touches text and the
        # form structure is preserved.
        # ============================================================
        pages_with_redactions: set[int] = set()
        redact_count = 0
        for translated_text, layout in pairs:
            page_idx = int((layout or {}).get("page", 0))
            bbox = (layout or {}).get("bbox")
            if bbox is None or page_idx >= len(doc):
                continue
            page = doc[page_idx]
            rect = fitz.Rect(*bbox)
            # Pad by 2pt (was 1pt) so descenders and kerning that fall
            # just outside the line bbox still get redacted. This is
            # the difference between a clean rebuild and visible
            # bilingual residue.
            mask_rect = fitz.Rect(
                rect.x0 - 2, rect.y0 - 2, rect.x1 + 2, rect.y1 + 2
            )
            try:
                page.add_redact_annot(mask_rect)
                pages_with_redactions.add(page_idx)
                redact_count += 1
            except Exception as e:
                logger.debug(
                    "add_redact_annot failed on page %d bbox %s: %s",
                    page_idx, bbox, e,
                )

        logger.info(
            "[pdf-rebuild v6] queued %d redactions across %d pages",
            redact_count, len(pages_with_redactions),
        )

        # Apply queued redactions per page with text-only semantics.
        # Older PyMuPDF builds may not have these enum constants, so we
        # fall back to the default if the kwargs are rejected.
        redaction_failures: dict[int, str] = {}
        for page_idx in pages_with_redactions:
            page = doc[page_idx]
            try:
                page.apply_redactions(
                    images=getattr(fitz, "PDF_REDACT_IMAGE_NONE", 0),
                    graphics=getattr(fitz, "PDF_REDACT_LINE_ART_NONE", 0),
                )
            except TypeError:
                try:
                    page.apply_redactions()
                except Exception as e:
                    redaction_failures[page_idx] = str(e)
                    logger.warning(
                        "apply_redactions failed on page %d: %s", page_idx, e
                    )
            except Exception as e:
                redaction_failures[page_idx] = str(e)
                logger.warning(
                    "apply_redactions failed on page %d: %s", page_idx, e
                )

        for translated_text, layout in pairs:
            page_idx = int((layout or {}).get("page", 0))
            bbox = (layout or {}).get("bbox")
            if bbox is None or page_idx >= len(doc):
                continue
            page = doc[page_idx]
            rect = fitz.Rect(*bbox)
            pix_for_page = page_pixmaps.get(page_idx)

            src = (layout or {}).get("source_text", "")
            font_size = float((layout or {}).get("font_size") or 11.0)

            # Noise blocks: redaction already removed the source text;
            # paint a small bg patch so any antialiased glyph residue
            # is covered, then move on.
            if _is_noise_text(src, font_size=font_size):
                bg_noise = _sample_background_color(
                    page, rect, pix=pix_for_page
                )
                page.draw_rect(
                    rect, color=None, fill=bg_noise, overlay=True
                )
                continue

            if not translated_text:
                continue

            # Sample the background colour and paint the bbox. If the
            # redaction for this page failed (some PDFs reject
            # apply_redactions for various reasons), paint a slightly
            # larger rectangle to physically hide the original glyphs
            # even though they're still in the text stream.
            bg = _sample_background_color(page, rect, pix=pix_for_page)
            if page_idx in redaction_failures:
                fail_safe = fitz.Rect(
                    rect.x0 - 2, rect.y0 - 2, rect.x1 + 2, rect.y1 + 2,
                )
                page.draw_rect(fail_safe, color=None, fill=bg, overlay=True)
            else:
                page.draw_rect(rect, color=None, fill=bg, overlay=True)

            color_int = int(layout.get("color") or 0)
            r = ((color_int >> 16) & 0xFF) / 255.0
            g = ((color_int >> 8) & 0xFF) / 255.0
            b = (color_int & 0xFF) / 255.0

            # ============================================
            # FONT SELECTION (style + script aware)
            # 1. If the target uses CJK, use the reserved CJK name we
            #    picked up-front (japan/korea/china-*).
            # 2. If the target uses Arabic / Hebrew / Hindi / Thai,
            #    register a system TTF for that script.
            # 3. Otherwise — pick the right Latin-script built-in
            #    (Times / Helvetica / Courier × Regular/Bold/Italic/
            #    Bold-Italic) using the ORIGINAL font name + flags so
            #    the translation visually matches the source.
            # ============================================
            flags = int((layout or {}).get("flags") or 0)
            src_font_name = (layout or {}).get("font_name") or ""

            if fontname != "helv":
                # CJK target → keep the CJK font name.
                page_font = fontname
            else:
                # Try a registered script TTF first.
                if page_idx not in registered_fonts:
                    registered_fonts[page_idx] = (
                        _register_script_font(page, target_lang) or ""
                    )
                script_font = registered_fonts[page_idx]
                if script_font:
                    page_font = script_font
                else:
                    # Latin-script fall-through: pick a style-matched
                    # PyMuPDF built-in so bold/italic/serif survive.
                    page_font = _map_pdf_font(src_font_name, flags)

            # RTL languages need reshaping + bidi reorder + right-align.
            display_text = (
                _shape_rtl(translated_text, target_lang)
                if is_rtl
                else translated_text
            )
            text_align = 2 if is_rtl else 0

            # ============================================
            # WIDTH-AWARE FONT SIZING
            # Measure the translation's rendered width at the source
            # font size. If it fits the bbox width, use the source
            # size verbatim — same typographic weight as the original.
            # If not, scale the size by (width_available / width_needed)
            # in one calculation rather than guessing through a ladder.
            # ============================================
            target_size = font_size
            try:
                rendered_w = fitz.get_text_length(
                    display_text, fontname=page_font, fontsize=font_size
                )
                box_w = max(1.0, rect.width - 2)  # small padding
                if rendered_w > box_w:
                    scale = box_w / rendered_w
                    target_size = max(6.0, font_size * scale)
            except Exception:
                pass

            # ============================================
            # DRAW
            # Use insert_textbox for proper line-wrapping inside the
            # bbox. If somehow it still overflows (very long
            # translation), fall through to a free-position insert_text
            # at the line's baseline so the segment isn't dropped.
            # ============================================
            drawn = False
            for size in (
                target_size,
                target_size * 0.95,
                target_size * 0.9,
                target_size * 0.85,
                max(target_size * 0.7, 6.0),
            ):
                rc = page.insert_textbox(
                    rect,
                    display_text,
                    fontsize=size,
                    color=(r, g, b),
                    fontname=page_font,
                    align=text_align,
                )
                if rc >= 0:
                    drawn = True
                    break

            if not drawn:
                try:
                    # Anchor at the line's baseline (approx 80% down
                    # the bbox for typical fonts) so the glyphs sit
                    # in the same vertical position as the original.
                    asc = float((layout or {}).get("ascender") or 0.8)
                    baseline_y = rect.y0 + asc * font_size
                    px = rect.x1 - 1 if is_rtl else rect.x0 + 1
                    page.insert_text(
                        fitz.Point(px, baseline_y),
                        display_text,
                        fontsize=max(target_size * 0.7, 6.0),
                        color=(r, g, b),
                        fontname=page_font,
                    )
                except Exception:
                    logger.debug(
                        f"Couldn't draw segment on page {page_idx} bbox {bbox}"
                    )

        doc.save(output_path, garbage=4, deflate=True)
    finally:
        doc.close()


def _build_bilingual_pdf(
    source_pdf_or_image: str,
    output_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
    source_is_image: bool,
) -> None:
    """Shared helper used by both image and PDF rebuilds.

    Produces a PDF with the source document on the first page(s) followed
    by a styled two-column translation table.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Image as RLImage,
        Table,
        TableStyle,
        PageBreak,
    )
    from reportlab.lib import colors

    page_w, page_h = A4

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=16,
        textColor=colors.HexColor("#1f2a2e"),
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#6b6558"),
        spaceAfter=14,
    )
    cell_style = ParagraphStyle(
        "Cell",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#1f2a2e"),
    )

    # ------------------------------------------------------------
    # Build the translation table on its own ReportLab PDF first.
    # We then merge with the source PDF (or wrap an image into a
    # PDF page) using PyMuPDF so the source pages stay pixel-perfect.
    # ------------------------------------------------------------
    table_buf = io.BytesIO()
    doc = SimpleDocTemplate(
        table_buf,
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    story: list = []
    story.append(Paragraph("Translation", title_style))
    story.append(
        Paragraph(
            "Each row matches one detected line in the source. Names, "
            "dates, ID numbers and reference codes are reproduced as-is.",
            subtitle_style,
        )
    )

    rows: list = [[Paragraph("<b>Source</b>", cell_style),
                   Paragraph("<b>Translation</b>", cell_style)]]
    for translated, layout in pairs:
        src = (layout or {}).get("source_text", "—") if layout else "—"
        if not (translated or "").strip() and not (src or "").strip():
            continue
        rows.append([
            Paragraph((src or "").replace("\n", "<br/>"), cell_style),
            Paragraph((translated or "").replace("\n", "<br/>"), cell_style),
        ])

    if len(rows) == 1:
        story.append(
            Paragraph("No text was detected in the source document.", cell_style)
        )
    else:
        col_w = (page_w - 30 * mm) / 2
        table = Table(rows, colWidths=[col_w, col_w], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f3ecdb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1f2a2e")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d8cfba")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.HexColor("#faf5ee")],
                    ),
                ]
            )
        )
        story.append(table)

    doc.build(story)
    table_buf.seek(0)

    # ------------------------------------------------------------
    # Compose the final PDF: source pages first, then the table pdf.
    # ------------------------------------------------------------
    import fitz

    out_pdf = fitz.open()

    # 1) Source page(s)
    if source_is_image:
        from PIL import Image
        pil = Image.open(source_pdf_or_image).convert("RGB")
        # Draw a small "Source document" header above the image.
        # Keeping it simple: just embed the image full-page.
        png_buf = io.BytesIO()
        pil.save(png_buf, format="PNG")
        png_buf.seek(0)
        pix = fitz.Pixmap(png_buf.getvalue())
        # Fit into A4 with a 15mm margin.
        margin = 15 * 72 / 25.4  # mm → pt
        page = out_pdf.new_page(width=page_w, height=page_h)
        avail_w = page_w - 2 * margin
        avail_h = page_h - 2 * margin
        scale = min(avail_w / pix.width, avail_h / pix.height)
        draw_w = pix.width * scale
        draw_h = pix.height * scale
        x = (page_w - draw_w) / 2
        y = (page_h - draw_h) / 2
        page.insert_image(
            fitz.Rect(x, y, x + draw_w, y + draw_h), pixmap=pix
        )
    else:
        # PDF source — embed every page as-is.
        try:
            src_pdf = fitz.open(source_pdf_or_image)
            out_pdf.insert_pdf(src_pdf)
            src_pdf.close()
        except Exception as e:
            logger.warning(f"Couldn't embed source PDF: {e}")

    # 2) Translation table page(s)
    table_pdf = fitz.open(stream=table_buf.getvalue(), filetype="pdf")
    out_pdf.insert_pdf(table_pdf)
    table_pdf.close()

    out_pdf.save(output_path, garbage=4, deflate=True)
    out_pdf.close()


def _rebuild_docx(
    original_path: str,
    output_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
) -> None:
    """Open the original DOCX, replace each paragraph/table-cell paragraph
    with its translation while keeping the run formatting (font, color,
    alignment) intact for visual similarity. Tables (boxes) survive
    because we never touch the table structure.
    """
    from docx import Document

    doc = Document(original_path)

    # Index the input pairs by their layout so we can look them up cheaply.
    by_para: dict[int, str] = {}
    by_cell: dict[tuple[int, int, int, int], str] = {}
    for translated, layout in pairs:
        kind = layout.get("kind")
        if kind == "docx_paragraph":
            by_para[int(layout.get("para_index", -1))] = translated
        elif kind == "docx_table_cell":
            by_cell[
                (
                    int(layout.get("table_index", -1)),
                    int(layout.get("row", -1)),
                    int(layout.get("col", -1)),
                    int(layout.get("para_in_cell", -1)),
                )
            ] = translated

    def _replace_paragraph_text(paragraph, new_text: str) -> None:
        """Keep the first run's formatting; drop any extra runs; set new text."""
        runs = paragraph.runs
        if not runs:
            paragraph.add_run(new_text)
            return
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""

    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx in by_para:
            _replace_paragraph_text(para, by_para[para_idx])

    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for cp_idx, cp in enumerate(cell.paragraphs):
                    key = (tbl_idx, row_idx, col_idx, cp_idx)
                    if key in by_cell:
                        _replace_paragraph_text(cp, by_cell[key])

    doc.save(output_path)


def _rebuild_image(
    original_path: str,
    output_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
    target_lang: str | None = None,
) -> None:
    """Produce a professional side-by-side translation PDF for ID/passport
    images.

    Why we don't overlay anymore
    ----------------------------
    OCR on documents with watermarks, holographic patterns, tightly packed
    rows or curved text returns noisy, overlapping bounding boxes. White-
    rectangle masking + inline redraw produced visible bleed-through of
    the source text, stacked translations, and unreadable output.

    Real human certified translators don't attempt visual overlay either
    — they produce a separate translation document. So we do the same:

    Page 1: the original image at full quality (visual reference / source
            of truth for the certifier).
    Page 2+: translated lines laid out cleanly in a two-column table —
             SOURCE on the left, TRANSLATION on the right. The lines
             appear in OCR order so the reader can align them visually
             with the original on page 1.
    """
    logger.info("[image-rebuild v6] high-quality overlay")
    from PIL import Image
    import fitz

    fontname = _pdf_font_for_language(target_lang)
    is_rtl = _is_rtl_language(target_lang)

    src = Image.open(original_path).convert("RGB")
    src_w, src_h = src.size
    img_buf = io.BytesIO()
    src.save(img_buf, format="PNG")
    img_buf.seek(0)
    src.close()

    out_pdf = fitz.open()
    page = out_pdf.new_page(width=src_w, height=src_h)
    pix = fitz.Pixmap(img_buf.getvalue())
    page.insert_image(page.rect, pixmap=pix)

    page_font = fontname
    if fontname == "helv":
        page_font = _register_script_font(page, target_lang) or "helv"

    # Pixmap of the source image used for both background-colour
    # sampling AND photo-region detection. We compute a 'colour
    # variance' for each candidate bbox — text on a flat background
    # has low variance, while text supposedly inside a holder's
    # photo or hologram has high variance. High-variance regions
    # are skipped to avoid painting nonsense over the photo.
    sample_pix = fitz.Pixmap(img_buf.getvalue())

    def _bbox_is_photographic(rect) -> bool:
        """True if the area inside `rect` looks like a photograph or
        hologram rather than flat ink-on-paper text. Heuristic: sample
        a grid of pixels and compute the variance of luminance values.
        High variance + wide RGB spread → photograph."""
        try:
            w = max(1, sample_pix.width)
            h = max(1, sample_pix.height)
            x0 = max(0, min(w - 1, int(rect.x0)))
            y0 = max(0, min(h - 1, int(rect.y0)))
            x1 = max(0, min(w - 1, int(rect.x1)))
            y1 = max(0, min(h - 1, int(rect.y1)))
            if x1 <= x0 + 2 or y1 <= y0 + 2:
                return False
            n_components = sample_pix.n
            samples = sample_pix.samples
            lums = []
            rs, gs, bs = [], [], []
            step_x = max(1, (x1 - x0) // 6)
            step_y = max(1, (y1 - y0) // 6)
            for yy in range(y0, y1, step_y):
                for xx in range(x0, x1, step_x):
                    idx = (yy * w + xx) * n_components
                    r = samples[idx]
                    g = samples[idx + 1]
                    b = samples[idx + 2]
                    rs.append(r); gs.append(g); bs.append(b)
                    lums.append(0.299 * r + 0.587 * g + 0.114 * b)
            if not lums:
                return False
            mean_l = sum(lums) / len(lums)
            var_l = sum((l - mean_l) ** 2 for l in lums) / len(lums)
            # Wide RGB spread within the bbox (photographs have varied
            # colour; printed text on white paper is mostly one colour
            # band). Threshold tuned empirically for ID cards.
            chroma = (max(rs) - min(rs)) + (max(gs) - min(gs)) + (max(bs) - min(bs))
            if var_l > 2000 and chroma > 200:
                return True
            return False
        except Exception:
            return False

    for translated, layout in pairs:
        bbox = (layout or {}).get("bbox")
        if not bbox:
            continue
        x0, y0, x1, y1 = bbox
        rect = fitz.Rect(x0, y0, x1, y1)

        src_text = (layout or {}).get("source_text", "")

        # Skip OCR boxes that fall in a photo region (holder's face,
        # hologram, watermark gradient). We'd produce garbage there.
        if _bbox_is_photographic(rect):
            continue

        bg = _sample_background_color(page, rect, pix=sample_pix)

        # Hide gibberish OCR blocks with the sampled background colour
        # so the patch blends with the document.
        if _is_noise_text(src_text):
            mask_rect = fitz.Rect(
                rect.x0 - 1, rect.y0 - 1, rect.x1 + 1, rect.y1 + 1
            )
            page.draw_rect(
                mask_rect, color=None, fill=bg, overlay=True
            )
            continue

        if not translated:
            continue

        # Prefer the font_size the extractor stored (median word
        # height, capped 7-12pt). Fall back to a bbox-derived
        # estimate for legacy projects.
        h = max(1.0, y1 - y0)
        font_size = float(
            (layout or {}).get("font_size") or max(6.0, h * 0.6)
        )
        font_size = min(font_size, 12.0)

        # IMAGE rebuild padding: bigger than PDF because OCR bbox
        # coordinates are inherently fuzzier than PyMuPDF's text-stream
        # bboxes. Pad horizontally by 25% of the bbox height and
        # vertically by 35% so the rectangle fully covers any source
        # glyph extending slightly outside the OCR-detected line.
        pad_v = max(2.0, h * 0.35)
        pad_h = max(2.0, h * 0.25)
        mask_rect = fitz.Rect(
            rect.x0 - pad_h,
            rect.y0 - pad_v,
            rect.x1 + pad_h,
            rect.y1 + pad_v,
        )
        page.draw_rect(mask_rect, color=None, fill=bg, overlay=True)

        display_text = _shape_rtl(translated, target_lang) if is_rtl else translated
        text_align = 2 if is_rtl else 0

        # 2. Try the derived font size, shrinking until it fits.
        drawn = False
        for size in (
            font_size,
            int(font_size * 0.95),
            int(font_size * 0.9),
            int(font_size * 0.85),
            int(font_size * 0.8),
            int(font_size * 0.75),
            int(font_size * 0.7),
            max(int(font_size * 0.6), 6),
        ):
            rc = page.insert_textbox(
                rect,
                display_text,
                fontsize=size,
                color=(0, 0, 0),
                fontname=page_font,
                align=text_align,
            )
            if rc >= 0:
                drawn = True
                break

        if not drawn:
            try:
                px = rect.x1 - 1 if is_rtl else rect.x0 + 1
                page.insert_text(
                    fitz.Point(px, rect.y0 + font_size),
                    display_text,
                    fontsize=max(int(font_size * 0.7), 6),
                    color=(0, 0, 0),
                    fontname=page_font,
                )
            except Exception:
                pass

    out_pdf.save(output_path, garbage=4, deflate=True)
    out_pdf.close()
    return
    # --- side-by-side renderer (kept for reference, unreachable) ---
    from PIL import Image
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Image as RLImage,
        Table,
        TableStyle,
        PageBreak,
    )
    from reportlab.lib import colors

    # --- Page 1: original image, scaled to fit A4 with margins ---
    pil = Image.open(original_path).convert("RGB")
    img_buf = io.BytesIO()
    pil.save(img_buf, format="PNG")
    img_buf.seek(0)

    page_w, page_h = A4
    max_w = page_w - 30 * mm
    max_h = page_h - 60 * mm  # leave space for header
    iw, ih = pil.size
    scale = min(max_w / iw, max_h / ih)
    rl_img = RLImage(img_buf, width=iw * scale, height=ih * scale)

    # Build PDF
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=16,
        textColor=colors.HexColor("#1f2a2e"),
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#6b6558"),
        spaceAfter=14,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#0a7870"),
        spaceAfter=10,
    )
    cell_style = ParagraphStyle(
        "Cell",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#1f2a2e"),
    )

    story: list = []
    story.append(Paragraph("Source document", title_style))
    story.append(
        Paragraph(
            "The image below is reproduced verbatim for reference. "
            "Translated content appears on the following pages.",
            subtitle_style,
        )
    )
    story.append(rl_img)
    story.append(PageBreak())

    # --- Page 2+: side-by-side translation table ---
    story.append(Paragraph("Translation", title_style))
    story.append(
        Paragraph(
            "Each row matches one detected line in the source. Names, "
            "dates, ID numbers and reference codes are reproduced as-is.",
            subtitle_style,
        )
    )

    rows: list = [[Paragraph("<b>Source</b>", cell_style),
                   Paragraph("<b>Translation</b>", cell_style)]]
    for translated, layout in pairs:
        # We don't have the source text here — pairs is (translated, layout).
        # The caller (translation_processor) will pass the source text in
        # layout["source_text"] so we can render it. Fall back to em-dash if
        # missing so the table still renders.
        src = (layout or {}).get("source_text", "—") if layout else "—"
        if not (translated or "").strip() and not (src or "").strip():
            continue
        rows.append([
            Paragraph((src or "").replace("\n", "<br/>"), cell_style),
            Paragraph((translated or "").replace("\n", "<br/>"), cell_style),
        ])

    if len(rows) == 1:
        # No translated lines at all — note it instead of an empty table.
        story.append(
            Paragraph(
                "No text was detected in the source image.",
                cell_style,
            )
        )
    else:
        col_w = (page_w - 30 * mm) / 2
        table = Table(rows, colWidths=[col_w, col_w], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f3ecdb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1f2a2e")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d8cfba")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.HexColor("#faf5ee")],
                    ),
                ]
            )
        )
        story.append(table)

    doc.build(story)


def _rebuild_txt(
    original_path: str,
    output_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
) -> None:
    """Stitch translated paragraphs back together with blank-line separators
    in the original paragraph order."""
    by_idx: dict[int, str] = {}
    for translated, layout in pairs:
        if layout.get("kind") == "txt_paragraph":
            by_idx[int(layout.get("para_index", -1))] = translated
    if not by_idx:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(t for t, _ in pairs))
        return
    ordered = [by_idx[k] for k in sorted(by_idx.keys())]
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(ordered))


def rebuild_output(
    source_kind: str,
    original_path: str,
    output_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
    target_lang: str | None = None,
) -> str:
    """Produce a file at `output_path` that resembles `original_path` but
    with the source text replaced by translations. Returns the path that
    was actually written (may differ from the requested extension when we
    fall back, e.g. an image source becomes a PDF)."""
    if source_kind == "PDF":
        _rebuild_pdf(original_path, output_path, pairs, target_lang=target_lang)
        return output_path
    if source_kind == "DOCX":
        _rebuild_docx(original_path, output_path, pairs)
        return output_path
    if source_kind == "IMAGE":
        # Force a .pdf extension since we render the result as a PDF.
        new_path = os.path.splitext(output_path)[0] + ".pdf"
        _rebuild_image(original_path, new_path, pairs, target_lang=target_lang)
        return new_path
    if source_kind == "TXT":
        _rebuild_txt(original_path, output_path, pairs)
        return output_path
    raise Exception(f"Don't know how to rebuild {source_kind}")
