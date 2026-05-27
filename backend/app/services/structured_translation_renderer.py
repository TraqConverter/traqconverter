"""Clean structured translation renderer.

The export flow now produces:

    Page 1            : the original document (image or PDF page) verbatim
    Page 2..N         : clean structured translation of the document body
                        following these formatting rules:
                          * preserve alignment (left / center / right)
                          * preserve bold / italic / all-caps emphasis
                          * preserve line spacing (single / extra-blank between
                            paragraph groups)
                          * non-text elements rendered as grey-italic
                            placeholders: [Coat of Arms], [Logo], [Photo],
                            [Stamp: text], [Signature], [Barcode], [QR Code],
                            [Illegible Text], [Handwritten text: …],
                            [Blanked Out]
                          * every line transcribed — no skipping
    Final page        : certification block (optional company logo at top)

The renderer accepts the same `pairs` list shape (translated_text, layout)
that the rest of the pipeline produces, plus the path to the original
document so page 1 can embed it verbatim.
"""
from __future__ import annotations

import io
import logging
import os
from typing import Any, Iterable

import fitz
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

logger = logging.getLogger(__name__)


# ============================================================
# Style palette — minimalist, matches Espresso aesthetic.
# ============================================================
_PALETTE = {
    "text": "#1f2a2e",
    "muted": "#6b6558",
    "rule": "#d8cfba",
    "placeholder": "#8a8270",  # grey for [Coat of Arms] etc.
    "accent": "#0a7870",
}


# Map "size" hints from Claude → point sizes that match a certified
# translator's transcription document. Tightened so body text reads
# at the same visual scale as the original (~9pt) — the previous
# 10.5pt made everything look oversized in the rebuild.
_SIZE_PT = {
    "small": 8.0,
    "normal": 9.0,
    "large": 11.0,
    "xlarge": 14.0,
}


# PyMuPDF / ReportLab built-in font families. We can't embed arbitrary
# TTFs without bundling them, so we pick the closest of the three
# Adobe-14 families. Claude classifies the source as one of these.
_FONT_TABLE = {
    "serif": {
        (False, False): "Times-Roman",
        (True, False):  "Times-Bold",
        (False, True):  "Times-Italic",
        (True, True):   "Times-BoldItalic",
    },
    "sans_serif": {
        (False, False): "Helvetica",
        (True, False):  "Helvetica-Bold",
        (False, True):  "Helvetica-Oblique",
        (True, True):   "Helvetica-BoldOblique",
    },
    "monospace": {
        (False, False): "Courier",
        (True, False):  "Courier-Bold",
        (False, True):  "Courier-Oblique",
        (True, True):   "Courier-BoldOblique",
    },
}


def _font_name(bold: bool, italic: bool, family: str = "serif") -> str:
    """Pick the right ReportLab built-in for the given family + style.

    Falls back to Times if the family is unknown.
    """
    fam = (family or "").lower()
    if fam not in _FONT_TABLE:
        fam = "serif"
    return _FONT_TABLE[fam][(bool(bold), bool(italic))]


_ALIGN = {"left": 0, "center": 1, "centre": 1, "right": 2, "justify": 4}


def _styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "placeholder": ParagraphStyle(
            "placeholder",
            parent=base["Italic"],
            fontName="Times-Italic",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor(_PALETTE["placeholder"]),
            spaceBefore=0,
            spaceAfter=0,
        ),
        "muted_small": ParagraphStyle(
            "muted_small",
            parent=base["Normal"],
            fontName="Times-Roman",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor(_PALETTE["muted"]),
        ),
        "cover_title": ParagraphStyle(
            "cover_title",
            parent=base["Heading1"],
            fontName="Times-Bold",
            fontSize=22,
            leading=28,
            textColor=colors.HexColor(_PALETTE["accent"]),
            alignment=1,
            spaceAfter=14,
        ),
        "cert_title": ParagraphStyle(
            "cert_title",
            parent=base["Heading1"],
            fontName="Times-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor(_PALETTE["text"]),
            alignment=1,
            spaceBefore=4,
            spaceAfter=12,
        ),
        "cert_body": ParagraphStyle(
            "cert_body",
            parent=base["Normal"],
            fontName="Times-Roman",
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor(_PALETTE["text"]),
            alignment=0,
        ),
    }


# ============================================================
# Helpers
# ============================================================

_BRACKET_PLACEHOLDER_RE = None  # imported lazily below to keep top clean


def _looks_like_placeholder(text: str) -> bool:
    """A line whose entire content is a `[ … ]` placeholder marker
    such as `[Coat of Arms]`, `[Stamp: COMUNE]`, `[Signature]`,
    `[Blanked Out]`, `[Illegible Text]`."""
    s = (text or "").strip()
    return len(s) >= 2 and s.startswith("[") and s.endswith("]")


def _alignment_from_layout(
    layout: dict[str, Any], page_width: float | None
) -> int:
    """Fallback alignment guess from bbox position when the extractor
    didn't supply an explicit 'alignment' value."""
    if not layout or not page_width:
        return 0
    bbox = layout.get("bbox")
    if not bbox or len(bbox) != 4:
        return 0
    try:
        x0, _, x1, _ = (float(v) for v in bbox)
    except Exception:
        return 0
    left_margin = x0 / page_width
    right_margin = (page_width - x1) / page_width
    centre_offset = abs(left_margin - right_margin)
    if centre_offset < 0.05 and left_margin > 0.1 and right_margin > 0.1:
        return 1  # center
    if right_margin > 0.5 and left_margin > 0.4:
        return 2  # right
    return 0


def _style_for_segment(
    text: str,
    rendered_text: str,
    layout: dict[str, Any],
    styles: dict[str, ParagraphStyle],
    page_width: float | None,
) -> tuple[ParagraphStyle, str]:
    """Pick the right style + escape the text for ReportLab.

    Uses the structured fields from Claude (alignment, bold, italic,
    all_caps, size hint) when present. Falls back to PDF span-flag
    bits + bbox-position guess for legacy segments.
    """
    body_text = (rendered_text or text or "").strip()

    # Placeholders always render in grey italic. Use the element's
    # font family if known so the placeholder visually belongs with
    # the surrounding body text (a sans-serif document gets
    # Helvetica-Oblique placeholders, etc.).
    placeholder_kind = (layout or {}).get("placeholder_kind")
    if placeholder_kind or _looks_like_placeholder(body_text):
        explicit_align = (layout or {}).get("alignment")
        aligned = _ALIGN.get(
            (explicit_align or "").lower(),
            _alignment_from_layout(layout, page_width),
        )
        family = (layout or {}).get("font_family") or ""
        ph_font = _font_name(False, True, family)
        ph_style = ParagraphStyle(
            f"ph_{ph_font}_{aligned}",
            fontName=ph_font,
            fontSize=10,
            leading=14,
            alignment=aligned,
            textColor=colors.HexColor(_PALETTE["placeholder"]),
        )
        safe = (
            body_text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        return ph_style, safe

    # Explicit structured fields take precedence.
    explicit_align = (layout or {}).get("alignment")
    aligned = _ALIGN.get(
        (explicit_align or "").lower(),
        _alignment_from_layout(layout, page_width),
    )

    is_bold = bool((layout or {}).get("bold"))
    is_italic = bool((layout or {}).get("italic"))
    is_all_caps = bool((layout or {}).get("all_caps"))

    # Legacy PDF segments don't have explicit bold/italic — read from
    # PyMuPDF span-flag bits.
    flags = int((layout or {}).get("flags") or 0)
    if not is_bold and (flags & 16):
        is_bold = True
    if not is_italic and (flags & 2):
        is_italic = True

    # Treat ALL CAPS source as bold (matches example: "COMUNE DI SAN
    # BONIFACIO" rendered in bold).
    src = (text or "").strip()
    if not is_all_caps and len(src) >= 4 and any(c.isalpha() for c in src):
        if src.upper() == src:
            is_all_caps = True
            is_bold = True

    # Claude classifies each element's font as serif / sans-serif /
    # monospace. The renderer picks the matching Adobe-14 family so
    # serif source documents render with Times, sans-serif with
    # Helvetica, and monospace certificate numbers etc. stay
    # monospaced.
    family = (layout or {}).get("font_family") or ""
    font_name = _font_name(is_bold, is_italic, family)

    # Size: explicit hint > stored font_size > default normal.
    size_hint = ((layout or {}).get("size_hint") or "").lower()
    if size_hint in _SIZE_PT:
        size = _SIZE_PT[size_hint]
    else:
        size = float((layout or {}).get("font_size") or 10.5)
        size = max(8.5, min(15.0, size))
    leading = size * 1.35

    derived = ParagraphStyle(
        f"derived_{font_name}_{size:.1f}_{aligned}",
        fontName=font_name,
        fontSize=size,
        leading=leading,
        alignment=aligned,
        textColor=colors.HexColor(_PALETTE["text"]),
        spaceBefore=0,
        spaceAfter=0,
    )
    safe = (
        body_text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    return derived, safe


# ============================================================
# Page 1 builder — original source embedded as the cover.
# ============================================================

def _add_original_pages(
    out_pdf: fitz.Document, original_path: str, source_kind: str
) -> None:
    """Embed the original document as the first page(s) of the export."""
    if source_kind == "PDF":
        try:
            src = fitz.open(original_path)
            out_pdf.insert_pdf(src)
            src.close()
            return
        except Exception as e:
            logger.warning("Couldn't embed source PDF as page 1: %s", e)

    if source_kind == "IMAGE":
        try:
            from PIL import Image
            pil = Image.open(original_path).convert("RGB")
            png_buf = io.BytesIO()
            pil.save(png_buf, format="PNG")
            png_buf.seek(0)
            pix = fitz.Pixmap(png_buf.getvalue())
            page = out_pdf.new_page(width=pix.width, height=pix.height)
            page.insert_image(page.rect, pixmap=pix)
            pil.close()
            return
        except Exception as e:
            logger.warning("Couldn't embed source image as page 1: %s", e)


# ============================================================
# Pages 2..N — clean structured translation.
# ============================================================

def _build_translation_pages(
    pairs: list[tuple[str, dict[str, Any]]],
    project_meta: dict[str, Any],
    company_logo_path: str | None,
) -> bytes:
    """Render the clean structured translation as a ReportLab PDF and
    return the bytes."""
    buf = io.BytesIO()
    styles = _styles()

    # Try to match source page size when available; default to A4.
    pw, ph = A4
    if project_meta.get("page_width") and project_meta.get("page_height"):
        try:
            pw = float(project_meta["page_width"])
            ph = float(project_meta["page_height"])
        except Exception:
            pass

    doc = SimpleDocTemplate(
        buf,
        pagesize=(pw, ph),
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=project_meta.get("title") or "Certified Translation",
    )

    story: list = []
    src_lang = project_meta.get("source_language") or ""
    tgt_lang = project_meta.get("target_language") or ""

    # The translated body starts directly here — no "CERTIFIED
    # TRANSLATION" cover header. The certified-statement block is
    # rendered on the final page instead.

    # Group pairs by page so the structured rendering visually follows
    # the original page breaks.
    last_page = None
    last_y_bottom: float | None = None
    for translated, layout in pairs:
        layout = layout or {}
        page_idx = int(layout.get("page", 0) or 0)
        if last_page is not None and page_idx != last_page:
            # Mirror source page break.
            story.append(PageBreak())
            last_y_bottom = None
        last_page = page_idx

        # Insert spacing that mirrors the gap between this line and the
        # previous one in the source. Treat vertical gaps > 1.5x line
        # height as a paragraph break.
        bbox = layout.get("bbox") or []
        gap_pt = 0.0
        if last_y_bottom is not None and len(bbox) == 4:
            try:
                gap_pt = max(0.0, float(bbox[1]) - last_y_bottom)
            except Exception:
                gap_pt = 0.0
        if gap_pt > 6:
            story.append(Spacer(1, min(gap_pt * 0.6, 18)))

        src_text = layout.get("source_text") or ""
        style, paragraph_text = _style_for_segment(
            src_text, translated, layout, styles, pw,
        )
        if not paragraph_text:
            continue
        story.append(Paragraph(paragraph_text, style))

        if len(bbox) == 4:
            try:
                last_y_bottom = float(bbox[3])
            except Exception:
                pass

    # Certification page — last page of the export.
    story.append(PageBreak())
    if company_logo_path and os.path.isfile(company_logo_path):
        try:
            # Keep the logo modestly sized (40mm wide) so it doesn't
            # dominate the page.
            img = RLImage(company_logo_path)
            iw, ih = img.imageWidth, img.imageHeight
            max_w = 50 * mm
            if iw > max_w:
                scale = max_w / iw
                img.drawWidth = iw * scale
                img.drawHeight = ih * scale
            story.append(img)
            story.append(Spacer(1, 8))
        except Exception as e:
            logger.warning("Couldn't embed company logo: %s", e)

    story.append(
        Paragraph("CERTIFIED TRANSLATION STATEMENT", styles["cert_title"])
    )
    cert_lines = project_meta.get("certification_lines") or []
    if not cert_lines:
        cert_lines = [
            "I hereby certify that the foregoing is a true and complete "
            "translation of the attached document.",
            "",
            f"Translator: {project_meta.get('translator_email', '')}",
            f"Date: {project_meta.get('certification_date', '')}",
            f"Source language: {src_lang or '—'}",
            f"Target language: {tgt_lang or '—'}",
        ]
    for line in cert_lines:
        if line:
            story.append(Paragraph(line, styles["cert_body"]))
        story.append(Spacer(1, 4))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# Public entry point
# ============================================================

def render_structured_export(
    *,
    source_kind: str,
    original_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
    project_meta: dict[str, Any],
    company_logo_path: str | None = None,
) -> bytes:
    """Produce the full export PDF as bytes.

    Layout: page 1+ = original document(s); pages N+ = clean structured
    translation; final page = certification block.
    """
    out_pdf = fitz.open()
    _add_original_pages(out_pdf, original_path, source_kind)

    translation_pdf_bytes = _build_translation_pages(
        pairs, project_meta, company_logo_path
    )
    try:
        translation_pdf = fitz.open(stream=translation_pdf_bytes, filetype="pdf")
        out_pdf.insert_pdf(translation_pdf)
        translation_pdf.close()
    except Exception as e:
        logger.warning("Couldn't append translation pages: %s", e)

    out = io.BytesIO()
    out_pdf.save(out)
    out_pdf.close()
    out.seek(0)
    return out.getvalue()


# ============================================================
# Structured DOCX renderer — same shape as the PDF output:
#   page 1     : original document (image embedded; PDFs are
#                rendered as page-sized images so they fit)
#   pages 2..N : structured translation with preserved alignment,
#                bold / italic, placeholders in grey italic
#   final page : certification block with optional company logo
# ============================================================

def _docx_alignment(value: int):
    """Map ReportLab alignment ints onto python-docx enum."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return {
        0: WD_ALIGN_PARAGRAPH.LEFT,
        1: WD_ALIGN_PARAGRAPH.CENTER,
        2: WD_ALIGN_PARAGRAPH.RIGHT,
        4: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(value, WD_ALIGN_PARAGRAPH.LEFT)


def _docx_font_name(family: str) -> str:
    """Map serif / sans_serif / monospace to a font installed almost
    everywhere — Word picks closest fallbacks on systems missing it."""
    fam = (family or "").lower()
    if fam == "sans_serif":
        return "Arial"
    if fam == "monospace":
        return "Courier New"
    return "Times New Roman"  # default + serif


def _docx_apply_paragraph_style(
    para,
    text: str,
    layout: dict[str, Any],
    is_placeholder: bool,
) -> None:
    """Configure a python-docx paragraph + its single run from the
    structured layout fields. Bold, italic, font family, size,
    alignment, and grey-italic placeholder colouring are honoured.
    """
    from docx.shared import Pt, RGBColor

    explicit_align = (layout or {}).get("alignment")
    aligned = _ALIGN.get(
        (explicit_align or "").lower(),
        _alignment_from_layout(layout, None),
    )
    para.alignment = _docx_alignment(aligned)

    # Compact spacing so the rebuild fits the same footprint as the
    # original (the default 8pt-after spacing on every paragraph
    # stretched the body across two pages).
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    is_bold = bool((layout or {}).get("bold"))
    is_italic = bool((layout or {}).get("italic"))
    is_all_caps = bool((layout or {}).get("all_caps"))
    flags = int((layout or {}).get("flags") or 0)
    if not is_bold and (flags & 16):
        is_bold = True
    if not is_italic and (flags & 2):
        is_italic = True

    # ALL-CAPS source maps to bold for visual hierarchy.
    if is_all_caps:
        is_bold = True

    size_hint = ((layout or {}).get("size_hint") or "").lower()
    if size_hint in _SIZE_PT:
        size_pt = _SIZE_PT[size_hint]
    else:
        size_pt = float((layout or {}).get("font_size") or 10.5)
        size_pt = max(8.5, min(15.0, size_pt))

    family = (layout or {}).get("font_family") or ""
    font_name = _docx_font_name(family)

    if is_placeholder:
        # Force italic, grey colour, regardless of source emphasis.
        run = para.add_run(text)
        run.italic = True
        run.font.name = font_name
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(
            int(_PALETTE["placeholder"][1:3], 16),
            int(_PALETTE["placeholder"][3:5], 16),
            int(_PALETTE["placeholder"][5:7], 16),
        )
        return

    run = para.add_run(text)
    run.bold = is_bold
    run.italic = is_italic
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(
        int(_PALETTE["text"][1:3], 16),
        int(_PALETTE["text"][3:5], 16),
        int(_PALETTE["text"][5:7], 16),
    )


def _embed_image_full_page(doc, image_path: str) -> None:
    """Drop an image as a single page in the DOCX, scaled to fit
    A4 portrait with reasonable margins. width=6 inches (instead of
    7) gives the image a safe top/bottom margin so it doesn't
    overflow to a second page (cause of the blank page 2 bug)."""
    from docx.shared import Inches
    p = doc.add_paragraph()
    run = p.add_run()
    try:
        run.add_picture(image_path, width=Inches(6))
    except Exception as e:
        logger.warning("DOCX: couldn't embed image %s: %s", image_path, e)
    doc.add_page_break()


def _embed_pdf_pages_as_images(doc, pdf_path: str) -> None:
    """For DOCX exports we render each source PDF page to a PNG and
    embed it page-by-page. width=6 inches keeps each page inside its
    own DOCX page (was 7, which let tall sources spill into a
    blank page after)."""
    from PIL import Image  # noqa: F401  (ensures PIL is importable)
    try:
        src = fitz.open(pdf_path)
        pages = list(src)
        for i, page in enumerate(pages):
            pix = page.get_pixmap(dpi=144)
            png_bytes = pix.tobytes("png")
            buf = io.BytesIO(png_bytes)
            buf.seek(0)
            from docx.shared import Inches
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(buf, width=Inches(6))
            # Page break between pages, AND before the structured
            # translation content. The trailing page break here is
            # what gives the translation its own clean page start.
            doc.add_page_break()
        src.close()
    except Exception as e:
        logger.warning("DOCX: couldn't embed source PDF pages: %s", e)


# ============================================================
# LAYOUT-AWARE DOCX RENDERER
# ------------------------------------------------------------
# Asks Claude to plan a DOCX-shaped JSON tree over the OCR
# elements, then emits paragraphs / rows / tables that mirror
# the visual structure of the original document. Falls through
# to the linear renderer on failure.
# ============================================================


def _set_cell_text(cell, eid, by_id, alignment_override=None, split_column=None):
    """Populate a table cell with a single styled paragraph from an
    element id. Uses the cell's existing default paragraph as the
    first content paragraph instead of removing it (removing the
    cell's last paragraph violates Word's OOXML invariant and Word
    silently drops every subsequent block in the document).

    `split_column` is "label" or "value" for form-table cells that
    share an id with the matching half. Splits on 2+ consecutive
    spaces between source_text's label and value.
    """
    if eid not in by_id:
        return
    translated, layout = by_id[eid]
    text = (translated or "").strip() or (layout.get("source_text") or "").strip()
    if not text:
        return

    # Form-table label/value split. The OCR emits things like
    # "Comune    SAN BENEDETTO DEL TRONTO" as a single segment; for
    # a layout-table column we want only the label or the value.
    if split_column in ("label", "value"):
        # Split on 2+ spaces or a tab — common form-row separator.
        import re as _re
        parts = _re.split(r" {2,}|\t+", text, maxsplit=1)
        if len(parts) == 2:
            text = parts[0].strip() if split_column == "label" else parts[1].strip()

    layout_for_render = dict(layout or {})
    if alignment_override in {"left", "center", "right", "justify"}:
        layout_for_render["alignment"] = alignment_override
    is_placeholder = bool(layout.get("placeholder_kind")) or (
        text.startswith("[") and text.endswith("]")
    )
    # Reuse the cell's default empty paragraph when it is still
    # empty; otherwise append a new one.
    para = None
    if cell.paragraphs and not cell.paragraphs[0].text:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    _docx_apply_paragraph_style(para, text, layout_for_render, is_placeholder)


def _set_cell_literal(cell, text, alignment="left"):
    """Variant of _set_cell_text for literal cell text (headers / spacers)."""
    if not text:
        return
    para = None
    if cell.paragraphs and not cell.paragraphs[0].text:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    _docx_apply_paragraph_style(para, text, {"alignment": alignment}, False)


def _ensure_trailing_paragraph(cell):
    """Word requires every <w:tc> to end with a <w:p>. When we add a
    nested table to a cell, Word treats the missing trailing
    paragraph as a corrupt document and stops rendering. This adds
    an empty paragraph if the cell's last block child is a table."""
    from docx.oxml.ns import qn

    tc = cell._element
    children = list(tc)
    # Find last block-level element (paragraph or table).
    for child in reversed(children):
        tag = child.tag
        if tag == qn("w:p"):
            return  # already ends with a paragraph
        if tag == qn("w:tbl"):
            break
        # Skip non-content elements like tcPr
    cell.add_paragraph("")


def _render_planned_block(
    doc_or_cell,
    block: dict[str, Any],
    by_id: dict[int, tuple[str, dict[str, Any]]],
    is_cell: bool = False,
) -> bool:
    """Recursively render one layout-planner block into a python-docx
    container (Document or _Cell). Returns True if anything was
    written. Each branch is wrapped in its own try/except so one
    malformed block can't black-hole the rest of the document.
    """
    from docx.shared import Pt

    try:
        kind = (block or {}).get("kind", "paragraph")

        if kind == "paragraph":
            try:
                eid = int(block.get("id"))
            except (TypeError, ValueError):
                return False
            if eid not in by_id:
                return False
            translated, layout = by_id[eid]
            text = (translated or "").strip() or (layout.get("source_text") or "").strip()
            if not text:
                return False
            layout_for_render = dict(layout or {})
            explicit_align = (block.get("alignment") or "").lower()
            if explicit_align in {"left", "center", "right", "justify"}:
                layout_for_render["alignment"] = explicit_align
            # When rendering into a cell, reuse the cell's default
            # empty paragraph for the FIRST paragraph so we don't end
            # up with an extra blank line above every cell.
            para = None
            if is_cell and doc_or_cell.paragraphs and not doc_or_cell.paragraphs[0].text:
                para = doc_or_cell.paragraphs[0]
            else:
                para = doc_or_cell.add_paragraph()
            is_placeholder = bool(layout.get("placeholder_kind")) or (
                text.startswith("[") and text.endswith("]")
            )
            _docx_apply_paragraph_style(para, text, layout_for_render, is_placeholder)
            if block.get("spacing_before_pt"):
                try:
                    para.paragraph_format.space_before = Pt(
                        float(block["spacing_before_pt"])
                    )
                except Exception:
                    pass
            return True

        if kind == "row":
            cols = block.get("columns") or []
            if not cols:
                return False
            tbl = doc_or_cell.add_table(rows=1, cols=len(cols))
            tbl.autofit = True
            for idx, col in enumerate(cols):
                cell = tbl.rows[0].cells[idx]
                col_blocks = col.get("blocks") or []
                if not col_blocks:
                    # Leave the cell's default empty paragraph alone.
                    continue
                # First child uses the cell's default paragraph; the
                # recursion knows to reuse it because is_cell=True.
                wrote_any = False
                for child in col_blocks:
                    if _render_planned_block(cell, child, by_id, is_cell=True):
                        wrote_any = True
                # Word requires the cell to end with a paragraph.
                _ensure_trailing_paragraph(cell)
            _strip_table_borders(tbl)
            return True

        if kind == "table":
            rows = block.get("rows") or []
            if not rows:
                return False
            cols_count = max((len(r) for r in rows), default=0)
            if cols_count == 0:
                return False
            tbl = doc_or_cell.add_table(rows=len(rows), cols=cols_count)
            tbl.autofit = True
            if block.get("border"):
                _apply_table_borders(tbl)
            else:
                _strip_table_borders(tbl)
            for ri, row in enumerate(rows):
                for ci in range(cols_count):
                    cell = tbl.rows[ri].cells[ci]
                    if ci >= len(row) or not row[ci]:
                        # Leave the cell's default empty paragraph.
                        continue
                    cellspec = row[ci]
                    if "id" in cellspec:
                        try:
                            eid = int(cellspec["id"])
                        except (TypeError, ValueError):
                            continue
                        _set_cell_text(
                            cell,
                            eid,
                            by_id,
                            alignment_override=(cellspec.get("alignment") or "").lower(),
                            split_column=cellspec.get("split_column"),
                        )
                    elif "text" in cellspec:
                        _set_cell_literal(
                            cell,
                            cellspec.get("text") or "",
                            alignment=(cellspec.get("alignment") or "left").lower(),
                        )
                    # Ensure cell ends with a paragraph.
                    _ensure_trailing_paragraph(cell)
            return True

        if kind == "spacer":
            from docx.shared import Pt as DocxPt
            para = doc_or_cell.add_paragraph()
            h_mm = float(block.get("height_mm") or 2.0)
            para.paragraph_format.space_after = DocxPt(h_mm * 2.83)
            return True

        return False
    except Exception as e:
        logger.warning("Block render failed (kind=%s): %s", block.get("kind"), e)
        return False


def _strip_table_borders(table) -> None:
    """Remove ALL borders from a python-docx Table so it looks like a
    pure layout grid rather than a visible table."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Replace any existing tblBorders with all-none.
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)


def _apply_table_borders(table) -> None:
    """Paint visible single-line borders on every edge of a table so
    parent/guardian tables and signature boxes actually look bordered
    in Word. python-docx tables have NO borders by default unless
    you apply a style or add OOXML borders explicitly."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Prefer the built-in Table Grid style when available — it gives
    # consistent borders even if the user's Word template overrides
    # tblBorders elsewhere.
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")          # 0.75pt line
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "1F2A2E")  # dark text color
        borders.append(b)
    tblPr.append(borders)


def _build_layout_plan_input(
    pairs: list[tuple[str, dict[str, Any]]],
) -> tuple[list[dict[str, Any]], dict[int, tuple[str, dict[str, Any]]], float, float]:
    """Build the (elements, by_id, page_w, page_h) tuple the planner
    needs. Element ids are assigned by index into the `pairs` list so
    the renderer can look the original pair back up after Claude
    returns a layout plan."""
    elements: list[dict[str, Any]] = []
    by_id: dict[int, tuple[str, dict[str, Any]]] = {}
    page_w = 0.0
    page_h = 0.0
    for idx, (translated, layout) in enumerate(pairs):
        layout = layout or {}
        bbox = layout.get("bbox") or [0, 0, 0, 0]
        try:
            x0, y0, x1, y1 = (float(v) for v in bbox)
        except Exception:
            x0, y0, x1, y1 = 0.0, 0.0, 0.0, 0.0
        page_w = max(page_w, x1)
        page_h = max(page_h, y1)
        elements.append(
            {
                "id": idx,
                "source": layout.get("source_text") or "",
                "translated": translated or "",
                "bbox": [x0, y0, x1, y1],
                "alignment": layout.get("alignment") or "left",
                "bold": bool(layout.get("bold")),
                "italic": bool(layout.get("italic")),
                "all_caps": bool(layout.get("all_caps")),
                "size": layout.get("size_hint") or layout.get("size") or "normal",
                "kind": "placeholder"
                if layout.get("placeholder_kind")
                else "text",
            }
        )
        by_id[idx] = (translated, layout)
    # Sensible defaults if everything was at 0,0.
    if page_w <= 1:
        page_w = 1000.0
    if page_h <= 1:
        page_h = 1400.0
    return elements, by_id, page_w, page_h


def render_planned_docx_export(
    *,
    source_kind: str,
    original_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
    project_meta: dict[str, Any],
    company_logo_path: str | None = None,
) -> bytes | None:
    """Layout-aware DOCX export. Returns the DOCX bytes, or None when
    the layout planner is unavailable or fails — callers should fall
    back to `render_structured_docx_export` in that case."""
    try:
        from app.services.layout_planner import plan_layout, is_available
    except Exception as e:
        logger.warning("Layout planner not importable: %s", e)
        return None

    if not is_available():
        return None

    elements, by_id, page_w, page_h = _build_layout_plan_input(pairs)
    if not elements:
        return None

    plan = plan_layout(elements, page_width=page_w, page_height=page_h)
    if not plan:
        return None

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Tighten section margins so the translated content has the same
    # available area as the original document (the linear renderer's
    # 1" margins left too little room and forced overflow onto a
    # second page).
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ---- Page 1+ : original document ----
    if source_kind == "IMAGE":
        _embed_image_full_page(doc, original_path)
    elif source_kind == "PDF":
        _embed_pdf_pages_as_images(doc, original_path)

    # ---- Translated section, structured per the layout plan ----
    for block in plan.get("blocks") or []:
        _render_planned_block(doc, block, by_id)

    # ---- Final page : certification (template or hardcoded) ----
    doc.add_page_break()
    _append_certification_page(
        doc,
        project_meta,
        company_logo_path,
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("Layout-aware DOCX renderer: produced %d-byte file", len(buf.getvalue()))
    return buf.getvalue()


def render_structured_docx_export(
    *,
    source_kind: str,
    original_path: str,
    pairs: list[tuple[str, dict[str, Any]]],
    project_meta: dict[str, Any],
    company_logo_path: str | None = None,
) -> bytes:
    """Produce the full export as a DOCX — same structural rules as
    the PDF builder above so the two exports look like siblings."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ---- Page 1+ : original document ----
    if source_kind == "IMAGE":
        _embed_image_full_page(doc, original_path)
    elif source_kind == "PDF":
        _embed_pdf_pages_as_images(doc, original_path)

    # ---- Pages N+ : structured translation ----
    src_lang = project_meta.get("source_language") or ""
    tgt_lang = project_meta.get("target_language") or ""

    last_page = None
    last_y_bottom: float | None = None
    for translated, layout in pairs:
        layout = layout or {}
        page_idx = int(layout.get("page", 0) or 0)
        if last_page is not None and page_idx != last_page:
            doc.add_page_break()
            last_y_bottom = None
        last_page = page_idx

        bbox = layout.get("bbox") or []
        if last_y_bottom is not None and len(bbox) == 4:
            try:
                gap_pt = max(0.0, float(bbox[1]) - last_y_bottom)
            except Exception:
                gap_pt = 0.0
            if gap_pt > 6:
                # Empty paragraph mirrors a paragraph-level gap.
                doc.add_paragraph("")

        body_text = (translated or "").strip()
        if not body_text:
            continue

        src_text = layout.get("source_text") or ""
        is_placeholder = bool(layout.get("placeholder_kind")) or (
            body_text.startswith("[") and body_text.endswith("]")
        )

        para = doc.add_paragraph()
        _docx_apply_paragraph_style(para, body_text, layout, is_placeholder)

        if len(bbox) == 4:
            try:
                last_y_bottom = float(bbox[3])
            except Exception:
                pass

    # ---- Final page : certification (template or hardcoded) ----
    doc.add_page_break()
    _append_certification_page(doc, project_meta, company_logo_path)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# CERT PAGE — template-aware
# ============================================================

def _append_certification_page(doc, project_meta, company_logo_path):
    """Append the final certification page to a python-docx Document.

    When the project has a cert template configured, the export
    pipeline pre-substitutes its {{tokens}} and stashes the bytes in
    `project_meta['certification_template_bytes']`. We copy every
    block of that template into the rebuild doc (preserving its
    formatting). Otherwise we render the previous hardcoded cert
    block, optionally prefixed by the user's logo.
    """
    from docx import Document as _Doc
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from copy import deepcopy
    import os as _os
    import io as _io

    template_bytes = project_meta.get("certification_template_bytes")
    if template_bytes:
        # User-provided template path. Walk every body element of the
        # substituted template and clone it into the rebuild doc so
        # the user's branding, layout, and inline images come through
        # intact.
        try:
            src_doc = _Doc(_io.BytesIO(template_bytes))
            for element in src_doc.element.body.iterchildren():
                # Skip the sectPr (section/page settings) — the
                # rebuild doc already has its own.
                if element.tag == qn("w:sectPr"):
                    continue
                doc.element.body.append(deepcopy(element))
            return
        except Exception as e:
            logger.warning(
                "Failed to embed cert template, falling back to "
                "default block: %s",
                e,
            )

    # Default hardcoded cert page.
    if company_logo_path and _os.path.isfile(company_logo_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(company_logo_path, width=Inches(2))
        except Exception as e:
            logger.warning("DOCX: couldn't embed logo: %s", e)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CERTIFIED TRANSLATION STATEMENT")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"
    title_run.font.color.rgb = RGBColor(
        int(_PALETTE["text"][1:3], 16),
        int(_PALETTE["text"][3:5], 16),
        int(_PALETTE["text"][5:7], 16),
    )

    src_lang = project_meta.get("source_language") or ""
    tgt_lang = project_meta.get("target_language") or ""
    cert_lines = project_meta.get("certification_lines") or [
        "I hereby certify that the foregoing is a true and complete "
        "translation of the attached document.",
        "",
        f"Translator: {project_meta.get('translator_email', '')}",
        f"Date: {project_meta.get('certification_date', '')}",
        f"Source language: {src_lang or '—'}",
        f"Target language: {tgt_lang or '—'}",
    ]
    for line in cert_lines:
        p = doc.add_paragraph()
        if line:
            r = p.add_run(line)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10.5)
